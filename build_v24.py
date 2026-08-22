"""
AutoDealership v24 — Complete Specification Document Builder
============================================================
Rebuilds entire document with:
1. Consistent styling (Inter font, unified colors)
2. All wireframes with full field spec tables + element logic
3. All APIs with complete request/response/error samples
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# DESIGN SYSTEM CONSTANTS
# ─────────────────────────────────────────────
FONT_BODY      = "Times New Roman"
FONT_CODE      = "Courier New"
FONT_HEADING   = "Times New Roman"

SZ_H1   = Pt(16)
SZ_H2   = Pt(13)
SZ_H3   = Pt(11)
SZ_BODY = Pt(10.5)
SZ_CODE = Pt(9)
SZ_TBL  = Pt(9.5)

# Colors (RGB)
C_H1_BG     = RGBColor(0x1D, 0x4E, 0xD8)   # Blue heading bg
C_H1_FG     = RGBColor(0xFF, 0xFF, 0xFF)   # White heading text
C_H2_FG     = RGBColor(0x1D, 0x4E, 0xD8)   # Blue
C_BODY      = RGBColor(0x1F, 0x29, 0x37)   # Dark gray body text
C_CODE_BG   = RGBColor(0xF3, 0xF4, 0xF6)   # Light gray code bg
C_TBL_HDR   = RGBColor(0x1D, 0x4E, 0xD8)   # Table header bg
C_TBL_HDR_F = RGBColor(0xFF, 0xFF, 0xFF)   # Table header text
C_TBL_ALT   = RGBColor(0xF0, 0xF4, 0xFF)   # Alternating row bg
C_SUCCESS   = RGBColor(0x06, 0x60, 0x2F)   # Green for success
C_ERROR     = RGBColor(0x99, 0x14, 0x0A)   # Red for error
C_WARN      = RGBColor(0x92, 0x40, 0x09)   # Orange for warning
C_NOTE      = RGBColor(0x1E, 0x40, 0xAF)   # Blue for notes

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb.red:02X}{rgb.green:02X}{rgb.blue:02X}')
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), f'{color.red:02X}{color.green:02X}{color.blue:02X}')
            tcBorders.append(el)
    tcPr.append(tcBorders)

def styled_run(para, text, font=FONT_BODY, size=SZ_BODY, color=C_BODY,
               bold=False, italic=False):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run

# ─────────────────────────────────────────────
# DOCUMENT HELPERS
# ─────────────────────────────────────────────

def add_h1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text.upper())
    run.font.name = FONT_HEADING
    run.font.size = SZ_H1
    run.font.color.rgb = C_H1_BG
    run.bold = True
    # Add bottom border line
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1D4ED8')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def add_h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = SZ_H2
    run.font.color.rgb = C_H2_FG
    run.bold = True
    return para

def add_h3(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = FONT_HEADING
    run.font.size = SZ_H3
    run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
    run.bold = True
    return para

def add_body(doc, text, color=C_BODY, bold=False, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = SZ_BODY
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.3 * (level + 1))
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = SZ_BODY
    run.font.color.rgb = C_BODY
    return para

def add_code_block(doc, code_text, label=None):
    """Add a styled code block with optional label"""
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(1)
        lr = lp.add_run(f"▶ {label}")
        lr.font.name = FONT_BODY
        lr.font.size = Pt(9)
        lr.font.color.rgb = C_H2_FG
        lr.bold = True

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(code_text)
    run.font.name = FONT_CODE
    run.font.size = SZ_CODE
    run.font.color.rgb = RGBColor(0x17, 0x20, 0x2E)
    # Shade background
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')
    pPr.append(shd)
    return para

def add_field_table(doc, rows_data):
    """
    Add a field spec table.
    rows_data: list of (field_id, component, required, validation, logic)
    """
    headers = ['Field ID', 'Component', 'Required', 'Validation Rule', 'Element Logic & Behavior']
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    widths = [Inches(1.1), Inches(1.0), Inches(0.7), Inches(1.6), Inches(3.1)]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = widths[i]

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_TBL_HDR)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = C_TBL_HDR_F
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_idx, row_data in enumerate(rows_data):
        row = table.add_row()
        if row_idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, C_TBL_ALT)
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT_BODY
            run.font.size = SZ_TBL
            run.font.color.rgb = C_BODY
            if i == 2 and val == 'Bắt buộc':
                run.font.color.rgb = C_ERROR
                run.bold = True

    doc.add_paragraph()
    return table

def add_api_spec(doc, method, endpoint, auth, desc,
                 request_body=None, path_params=None, query_params=None,
                 happy_response=None, error_responses=None):
    """Add a complete API specification block"""
    # Method + Endpoint header
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)

    # Method badge color
    method_colors = {
        'GET': RGBColor(0x06, 0x60, 0x2F),
        'POST': RGBColor(0x1D, 0x4E, 0xD8),
        'PUT': RGBColor(0x92, 0x40, 0x09),
        'PATCH': RGBColor(0x6B, 0x21, 0xA8),
        'DELETE': RGBColor(0x99, 0x14, 0x0A),
    }
    mc = method_colors.get(method, C_BODY)

    r1 = para.add_run(f" {method} ")
    r1.font.name = FONT_CODE
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.bold = True

    # Shade method badge
    # Add method as bold colored text
    r2 = para.add_run(f"  {endpoint}")
    r2.font.name = FONT_CODE
    r2.font.size = Pt(10)
    r2.font.color.rgb = mc
    r2.bold = True

    # Auth + description
    if auth or desc:
        info_para = doc.add_paragraph()
        info_para.paragraph_format.space_after = Pt(2)
        info_para.paragraph_format.left_indent = Inches(0.2)
        if auth:
            ar = info_para.add_run(f"Auth: {auth}  ")
            ar.font.name = FONT_BODY
            ar.font.size = Pt(9)
            ar.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            ar.bold = True
        if desc:
            dr = info_para.add_run(desc)
            dr.font.name = FONT_BODY
            dr.font.size = Pt(9)
            dr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Path params
    if path_params:
        add_code_block(doc, path_params, label="Path Parameters")

    # Query params
    if query_params:
        add_code_block(doc, query_params, label="Query Parameters")

    # Request body
    if request_body:
        add_code_block(doc, request_body, label="Request Body (JSON)")

    # Happy path response
    if happy_response:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(1)
        lr = lp.add_run("✅ Success Response")
        lr.font.name = FONT_BODY
        lr.font.size = Pt(9)
        lr.font.color.rgb = C_SUCCESS
        lr.bold = True
        add_code_block(doc, happy_response)

    # Error responses
    if error_responses:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(1)
        lr = lp.add_run("❌ Error Responses")
        lr.font.name = FONT_BODY
        lr.font.size = Pt(9)
        lr.font.color.rgb = C_ERROR
        lr.bold = True
        add_code_block(doc, error_responses)

    # Divider
    dp = doc.add_paragraph()
    dp.paragraph_format.space_after = Pt(4)
    dr = dp.add_run("─" * 80)
    dr.font.name = FONT_CODE
    dr.font.size = Pt(7)
    dr.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)

def add_wireframe(doc, ascii_art):
    """Add wireframe as monospaced code block"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.1)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(ascii_art)
    run.font.name = FONT_CODE
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x17, 0x20, 0x2E)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F8FAFC')
    pPr.append(shd)

def add_state_machine_table(doc, rows):
    headers = ['From Status', 'To Status', 'Trigger / Event', 'Guard Condition', 'Actor']
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_TBL_HDR)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = C_TBL_HDR_F
        run.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT_BODY
            run.font.size = SZ_TBL
            run.font.color.rgb = C_BODY
    doc.add_paragraph()
    return table

def add_notif_table(doc, rows):
    headers = ['Event', 'Channels', 'Recipient', 'Template', 'Timing']
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, C_TBL_HDR)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        run.font.color.rgb = C_TBL_HDR_F
        run.bold = True
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, C_TBL_ALT)
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT_BODY
            run.font.size = SZ_TBL
            run.font.color.rgb = C_BODY
    doc.add_paragraph()

def add_generic_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, C_TBL_HDR)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.font.name = FONT_BODY
        r.font.size = Pt(9)
        r.font.color.rgb = C_TBL_HDR_F
        r.bold = True
    for idx, row_data in enumerate(rows):
        row = table.add_row()
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_bg(cell, C_TBL_ALT)
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(val))
            r.font.name = FONT_BODY
            r.font.size = SZ_TBL
            r.font.color.rgb = C_BODY
    doc.add_paragraph()
    return table

# ═══════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# Set default font
doc.styles['Normal'].font.name = FONT_BODY
doc.styles['Normal'].font.size = SZ_BODY

# ───────────────────────────────────
# COVER / TITLE
# ───────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(20)
tr = title_para.add_run('TÀI LIỆU ĐẶC TẢ HỆ THỐNG TOÀN DIỆN\nAutoDealership Enterprise Platform')
tr.font.name = FONT_HEADING
tr.font.size = Pt(20)
tr.font.color.rgb = C_H1_BG
tr.bold = True

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_para.add_run('Phiên bản 24.0 — Master Final Enterprise\n(Đồng Bộ Review Findings + Full API Spec + Full Wireframe Field Tables)\nCập nhật: 2026-08-23')
sr.font.name = FONT_BODY
sr.font.size = Pt(10)
sr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

doc.add_page_break()

# ───────────────────────────────────
# SECTION 1: ROLE-BASED WORKFLOWS
# ───────────────────────────────────
add_h1(doc, '1. Sơ Đồ High-Level Phân Luồng Tính Năng Theo Vai Trò (RBAC Workflows)')

add_body(doc,
    'Toàn bộ đội ngũ Frontend, Backend và QA Automation nắm bức tranh tổng thể trước khi '
    'đi vào chi tiết từng màn hình. Hệ thống có 4 vai trò chính: CUSTOMER, SALE, MANAGER, ADMIN.')

add_h2(doc, '1.1. Luồng Khách Hàng (Customer Journey)')
add_bullet(doc, 'Đăng ký tài khoản (SCR-00-REG) → Xác thực OTP → Hoàn thiện hồ sơ (SCR-00-PROF)')
add_bullet(doc, 'Khám phá danh mục xe (SCR-01) → Xem chi tiết PDP (SCR-02)')
add_bullet(doc, 'Đặt lịch lái thử (SCR-03) → Trải nghiệm xe demo')
add_bullet(doc, 'Đặt cọc thanh toán (SCR-04 → SCR-10) → Theo dõi đơn hàng (SCR-11)')
add_bullet(doc, 'Mua trả góp: nộp hồ sơ vay → Đổi ngân hàng nếu bị từ chối (SCR-11)')
add_bullet(doc, 'Thu cũ đổi mới: Định giá xe cũ → Cấn trừ dòng tiền (SCR-02 → SCR-04)')
add_bullet(doc, 'Hoàn cọc nếu cần (SCR-11 → SCR-07)')

add_h2(doc, '1.2. Luồng Nhân Viên Kinh Doanh (Sale Consultant)')
add_bullet(doc, 'Quản lý phễu CRM Leads Kanban (SCR-06) → Gọi 1-click → Gửi PayLink')
add_bullet(doc, 'Tạo đơn hộ khách hàng (SCR-04), hỗ trợ trade-in, hỗ trợ refund')
add_bullet(doc, 'Xem kho xe available, phối hợp Manager để Hold VIN (SCR-08)')

add_h2(doc, '1.3. Luồng Cửa Hàng Trưởng (Store Manager)')
add_bullet(doc, 'Duyệt chiết khấu cấp Showroom (SCR-07)')
add_bullet(doc, 'Quản lý kho xe, tạo Hold VIN 24h, điều chuyển kho (SCR-08 → SCR-12)')
add_bullet(doc, 'Manager Override refund khi thiếu chứng từ (SCR-07)')
add_bullet(doc, 'Xem KPI Dashboard báo cáo doanh thu Showroom (SCR-05)')

add_h2(doc, '1.4. Luồng Quản Trị Viên (Super Admin)')
add_bullet(doc, 'Quản trị Users và RBAC Matrix (SCR-09)')
add_bullet(doc, 'Cấu hình tham số hệ thống (SCR-09-CONF) — dùng bảng system_configs')
add_bullet(doc, 'Duyệt chiết khấu tích lũy > 30M (SCR-07)')
add_bullet(doc, 'Xem Audit Log + Decrypt PII (SCR-13)')

add_h2(doc, '1.5. Ma Trận Phân Quyền Cấp Màn Hình (Screen-Level RBAC)')
add_generic_table(doc,
    ['Màn hình', 'Chức năng', 'Admin', 'Manager', 'Sale', 'Customer'],
    [
        ['SCR-00', 'Đăng nhập / Đăng xuất', '✅', '✅', '✅', '✅'],
        ['SCR-00-REG', 'Đăng ký tài khoản', 'N/A', 'N/A', 'N/A', '✅'],
        ['SCR-00-FP', 'Quên mật khẩu', '✅', '✅', '✅', '✅'],
        ['SCR-00-PROF', 'Hồ sơ cá nhân', '✅', '✅', '✅', '✅'],
        ['SCR-01', 'Danh mục xe', '✅', '✅', '✅', '✅'],
        ['SCR-02', 'Chi tiết xe (PDP)', '✅', '✅', '✅', '✅'],
        ['SCR-03', 'Đặt lịch lái thử', '✅', '✅', '✅ (hộ)', '✅'],
        ['SCR-04', 'Checkout đặt cọc', 'Hỗ trợ', 'Hỗ trợ', 'Tạo hộ', '✅'],
        ['SCR-05', 'Dashboard KPI', 'All-Margin', 'Showroom', 'Cá nhân', '—'],
        ['SCR-06', 'CRM Leads Kanban', 'Báo cáo', 'Toàn showroom', 'Lead cá nhân', '—'],
        ['SCR-07', 'Duyệt chiết khấu', 'Admin level', 'Showroom level', 'Tạo request', '—'],
        ['SCR-08', 'Kho xe & Quota', 'Full CRUD', 'Hold VIN, chuyển kho', 'Xem AVAILABLE', '—'],
        ['SCR-09', 'Quản trị Users', 'Full CRUD', '—', '—', '—'],
        ['SCR-09-CONF', 'Cấu hình hệ thống', 'Full CRUD', '—', '—', '—'],
        ['SCR-10', 'Kết quả thanh toán', '—', '—', 'Tạo lại link', '✅ Xem'],
        ['SCR-11', 'Cổng đơn hàng', '—', '—', '—', '✅ Xem & thao tác'],
        ['SCR-12', 'Điều chuyển kho', 'Full', 'Tạo/Duyệt', '—', '—'],
        ['SCR-13', 'Audit logs', 'Full + Decrypt', '—', '—', '—'],
    ]
)

doc.add_page_break()

# ───────────────────────────────────
# SECTION 2: SCREEN FLOWS
# ───────────────────────────────────
add_h1(doc, '2. Đặc Tả Chi Tiết Các Luồng Màn Hình')

# ─── SCR-00: Login ───
add_h2(doc, '2.1. Luồng 1: Xác Thực & Đăng Nhập (SCR-00, SCR-00-REG, SCR-00-FP, SCR-00-PROF)')

add_h3(doc, '2.1.1. SCR-00 — Đăng Nhập / Đăng Xuất / Mock OTP')
add_bullet(doc, 'URL: /auth/login  |  Prerequisites: Chưa đăng nhập')
add_bullet(doc, 'Navigation: Entry point. [Quên mật khẩu] → SCR-00-FP. [Đăng ký] → SCR-00-REG. Login thành công → /home')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [🚗 AUTO DEALERSHIP]                          Hotline: 1900-xxxx        │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│                                                                          │\n'
'│  ┌─────────────────────────────────────────────┐                        │\n'
'│  │           ĐĂNG NHẬP TÀI KHOẢN              │                        │\n'
'│  │                                             │                        │\n'
'│  │  [  Mật khẩu  ]    [  OTP  ]   ← tab_auth_mode                    │\n'
'│  │                                             │                        │\n'
'│  │  Email hoặc Số điện thoại *                 │                        │\n'
'│  │  [ identity_input                      ]   │                        │\n'
'│  │                                             │                        │\n'
'│  │  Mật khẩu *                         [👁]   │  ← tab: PASSWORD       │\n'
'│  │  [ password                            ]   │                        │\n'
'│  │  Sai 5 lần → khóa 30 phút                  │                        │\n'
'│  │                                             │                        │\n'
'│  │  [       Quên mật khẩu?      ]             │                        │\n'
'│  │                                             │                        │\n'
'│  │  [           ĐĂNG NHẬP           ]  🔵    │                        │\n'
'│  │                                             │                        │\n'
'│  │  ─────────── hoặc ──────────────           │                        │\n'
'│  │                                             │                        │\n'
'│  │  [       Chưa có tài khoản? Đăng ký  ]     │                        │\n'
'│  └─────────────────────────────────────────────┘                        │\n'
'│                                                                          │\n'
'│  Tab OTP:                                                                │\n'
'│  Số điện thoại: [ otp_phone_input    ]  [Gửi OTP]                       │\n'
'│  Mã OTP: [ ][ ][ ][ ][ ][ ]  ⏱ 60s   ← otp_digits (888888 sandbox)    │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-00')
add_field_table(doc, [
    ('tab_auth_mode',   'Tab/Toggle',         'Bắt buộc', 'Enum: PASSWORD | OTP',
     'Chuyển đổi giao diện. Default: PASSWORD. Lưu preference vào localStorage. Khi chuyển tab → reset form fields.'),
    ('identity_input',  'Text Field',         'Bắt buộc', '^(0|84)[35789][0-9]{8}$ hoặc email RFC5322',
     'Nhận email (admin@autodealer.vn) hoặc SĐT VN. Auto-detect format để route sang đúng auth flow. Trim whitespace.'),
    ('password',        'Password Field',     'Bắt buộc (tab PW)', 'Min 10, 1 hoa, 1 thường, 1 số, 1 ký tự đặc biệt',
     'Icon 👁 toggle show/hide. Đếm failed_login_attempts: sau 5 lần sai → lock 30 phút, hiện countdown timer. Không gửi nếu field trống.'),
    ('btn_forgot_pw',   'Link Button',        'Tùy chọn', 'N/A',
     'Điều hướng → /auth/forgot-password. Không cần confirm.'),
    ('btn_login',       'Primary Button',     'N/A', 'N/A',
     'Disabled nếu identity_input rỗng hoặc (PW tab: password rỗng) hoặc (OTP tab: otp_digits chưa đủ 6 số). Loading spinner khi đang call API.'),
    ('btn_register',    'Secondary Link',     'N/A', 'N/A',
     'Điều hướng → /auth/register.'),
    ('otp_phone_input', 'Text Field',         'Bắt buộc (tab OTP)', '^(0|84)[35789][0-9]{8}$',
     'SĐT nhận OTP. Sau khi nhập đúng format → enable [Gửi OTP].'),
    ('btn_send_otp',    'Secondary Button',   'N/A', 'N/A',
     'Gọi POST /auth/send-otp. Sau khi gửi → disable 60s countdown. Gửi tối đa 3 lần/phiên.'),
    ('otp_digits',      '6-box Input',        'Bắt buộc (tab OTP)', '6 chữ số. Sandbox: 888888. Production: random TTL 60s',
     'Auto-focus ô tiếp theo khi nhập xong. Auto-submit khi đủ 6 số. Đếm ngược 60s từ lúc gửi. Sai 3 lần → hủy phiên OTP.'),
])

# API SCR-00
add_h3(doc, 'API Specs — SCR-00')

add_api_spec(doc, 'POST', '/api/v1/auth/send-otp',
    auth='Public. Rate limit: 3 req/phone/hour.',
    desc='Gửi mã OTP 6 số về SĐT. Sandbox luôn dùng 888888.',
    request_body='{\n  "phone": "0901234567"    // required, VN phone format\n}',
    happy_response='HTTP 200 OK\n{\n  "success": true,\n  "data": {\n    "message": "OTP đã được gửi tới 090*****67",\n    "expires_in": 60,\n    "resend_after": 60\n  }\n}',
    error_responses=(
        'HTTP 400 Bad Request\n{ "success": false, "error": { "code": "ERR_INVALID_PHONE", "message": "Số điện thoại không hợp lệ" } }\n\n'
        'HTTP 429 Too Many Requests\n{ "success": false, "error": { "code": "ERR_OTP_RATE_LIMIT", "message": "Bạn đã yêu cầu OTP quá nhiều lần. Thử lại sau 1 giờ." } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/auth/login',
    auth='Public.',
    desc='Đăng nhập bằng mật khẩu hoặc OTP. Trả về JWT access + refresh token.',
    request_body=(
        '// Đăng nhập mật khẩu:\n'
        '{\n'
        '  "tab_mode": "PASSWORD",\n'
        '  "identity": "user@example.com",  // email hoặc phone\n'
        '  "password": "Abc@123456"\n'
        '}\n\n'
        '// Đăng nhập OTP:\n'
        '{\n'
        '  "tab_mode": "OTP",\n'
        '  "identity": "0901234567",\n'
        '  "otp_code": "888888"\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "access_token": "eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9...",  // TTL: 15 phút\n'
        '    "refresh_token": "eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9...", // TTL: 7 ngày, HttpOnly cookie\n'
        '    "user": {\n'
        '      "id": "550e8400-e29b-41d4-a716-446655440000",\n'
        '      "email": "user@example.com",\n'
        '      "full_name": "Nguyen Van A",\n'
        '      "role": "CUSTOMER",\n'
        '      "showroom_id": null\n'
        '    }\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 401 Unauthorized\n{ "success": false, "error": { "code": "ERR_AUTH_001", "message": "Sai email/SĐT hoặc mật khẩu" } }\n\n'
        'HTTP 423 Locked\n{ "success": false, "error": { "code": "ERR_AUTH_002", "message": "Tài khoản bị khóa 30 phút. Còn lại: 25 phút." } }\n\n'
        'HTTP 400 Bad Request (OTP expired)\n{ "success": false, "error": { "code": "ERR_OTP_EXPIRED", "message": "Mã OTP đã hết hạn. Vui lòng yêu cầu mã mới." } }\n\n'
        'HTTP 400 Bad Request (OTP invalid)\n{ "success": false, "error": { "code": "ERR_OTP_INVALID", "message": "Mã OTP không đúng. Còn 2 lần thử." } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/auth/logout',
    auth='Bearer token.',
    desc='Đăng xuất. Invalidate refresh token hiện tại, xóa session khỏi user_sessions.',
    request_body='{\n  "refresh_token": "eyJhbGc..."  // optional nếu dùng HttpOnly cookie\n}',
    happy_response='HTTP 200 OK\n{\n  "success": true,\n  "data": { "message": "Đăng xuất thành công" }\n}',
    error_responses='HTTP 401 Unauthorized\n{ "success": false, "error": { "code": "ERR_AUTH_INVALID_TOKEN", "message": "Token không hợp lệ" } }'
)

add_api_spec(doc, 'POST', '/api/v1/auth/refresh-token',
    auth='Refresh token (HttpOnly cookie hoặc body).',
    desc='Lấy access token mới. Rotation: phát refresh token mới, invalidate cái cũ.',
    request_body='{\n  "refresh_token": "eyJhbGc..."  // nếu không dùng cookie\n}',
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "access_token": "eyJhbGciOiJSUzI1NiIs...",\n'
        '    "refresh_token": "eyJhbGciOiJSUzI1NiIs..." // token mới (rotation)\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 401 Unauthorized\n{ "success": false, "error": { "code": "ERR_REFRESH_INVALID", "message": "Refresh token không hợp lệ hoặc đã hết hạn" } }\n\n'
        'HTTP 401 (token reuse detected)\n{ "success": false, "error": { "code": "ERR_TOKEN_REUSE", "message": "Phát hiện tái sử dụng token. Tất cả phiên đã bị thu hồi." } }'
    )
)

# ─── SCR-00-REG ───
add_h3(doc, '2.1.2. SCR-00-REG — Đăng Ký Tài Khoản Mới')
add_bullet(doc, 'URL: /auth/register  |  Prerequisites: Chưa có tài khoản')
add_bullet(doc, 'Navigation: Từ SCR-00 [Đăng ký]. Sau đăng ký thành công → tự đăng nhập → /home')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [🚗 AUTO DEALERSHIP]                          ĐĂNG KÝ TÀI KHOẢN        │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│                                                                          │\n'
'│  Họ và tên *                                                             │\n'
'│  [ full_name                                              ]              │\n'
'│                                                                          │\n'
'│  Email *                                                                 │\n'
'│  [ email                                                  ]              │\n'
'│  ⓘ Bạn sẽ nhận link xác thực sau khi đăng ký                           │\n'
'│                                                                          │\n'
'│  Số điện thoại *                                                         │\n'
'│  [ phone (VN: 09x / 03x / 07x / 08x)                     ]              │\n'
'│                                                                          │\n'
'│  Mật khẩu *                                           [👁]               │\n'
'│  [ password                                               ]              │\n'
'│  ⓘ Min 10 ký tự, 1 hoa, 1 thường, 1 số, 1 ký tự đặc biệt (@$!%*?&)   │\n'
'│  Strength: [████████░░] Strong                                           │\n'
'│                                                                          │\n'
'│  Xác nhận mật khẩu *                                  [👁]               │\n'
'│  [ confirm_password                                       ]              │\n'
'│                                                                          │\n'
'│  Mã OTP (gửi về SĐT) *    [  Nhận mã OTP  ]                             │\n'
'│  [ ][ ][ ][ ][ ][ ]   ⏱ Còn 60s                                        │\n'
'│                                                                          │\n'
'│  ☐ Tôi đồng ý với Điều khoản sử dụng và Chính sách bảo mật             │\n'
'│                                                                          │\n'
'│  [         Đã có tài khoản? Đăng nhập         ]                         │\n'
'│  [                    ĐĂNG KÝ                     ] 🔵 (disabled trước) │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-00-REG')
add_field_table(doc, [
    ('full_name',        'Text Field',       'Bắt buộc', 'Min 2, Max 100 ký tự. Không chứa ký tự đặc biệt.',
     'Hiển thị trên UI header sau login. Dùng in hợp đồng. Trim whitespace hai đầu.'),
    ('email',            'Text Field',       'Bắt buộc', 'RFC 5322 format. UNIQUE trong hệ thống.',
     'Sau đăng ký: gửi verification email. email_verified=false cho đến khi click link. Lowercase tự động.'),
    ('phone',            'Text Field',       'Bắt buộc', '^(0|84)[35789][0-9]{8}$. UNIQUE.',
     'Dùng để đăng nhập OTP. Nhận SMS brandname. Auto format khi nhập 84xxx → 0xxx.'),
    ('password',         'Password Field',   'Bắt buộc', '^(?=.*[A-Z])(?=.*[a-z])(?=.*\\d)(?=.*[@$!%*?&])[A-Za-z\\d@$!%*?&]{10,}$',
     'Password strength indicator (4 levels: Weak/Fair/Good/Strong). Hash bcrypt cost=12. Icon toggle show/hide.'),
    ('confirm_password', 'Password Field',   'Bắt buộc', 'Phải khớp chính xác với password.',
     'Validate client-side only, KHÔNG gửi lên server. Hiện lỗi ngay khi blur nếu không khớp.'),
    ('btn_send_otp',     'Secondary Button', 'N/A', 'N/A',
     'Disabled nếu phone chưa hợp lệ. Sau click → POST /auth/send-otp. Cooldown 60s sau khi gửi. Max 3 lần gửi.'),
    ('otp_digits',       '6-box Input',      'Bắt buộc', '6 chữ số. Sandbox: 888888. TTL 60s.',
     'Auto-focus tiếp theo. Auto-paste khi copy từ SMS. Countdown từ 60s → 0 → [Gửi lại mã].'),
    ('terms_checkbox',   'Checkbox',         'Bắt buộc', 'Must be checked = true.',
     'Nút [Đăng ký] disabled nếu chưa tích. Link "Điều khoản" → mở modal hoặc tab mới.'),
    ('btn_register',     'Primary Button',   'N/A', 'N/A',
     'Disabled cho đến khi tất cả required fields hợp lệ VÀ terms_checkbox=true. Loading spinner khi submit. Gọi POST /auth/register.'),
    ('link_login',       'Text Link',        'N/A', 'N/A',
     'Điều hướng → /auth/login. Không hiện confirm dialog.'),
])

add_h3(doc, 'API Specs — SCR-00-REG')

add_api_spec(doc, 'POST', '/api/v1/auth/register',
    auth='Public. Rate limit: 5 req/IP/hour.',
    desc='Đăng ký tài khoản khách hàng mới. Tạo user + customer_profile cơ bản.',
    request_body=(
        '{\n'
        '  "full_name": "Nguyen Van A",\n'
        '  "email": "nguyenvana@gmail.com",\n'
        '  "phone": "0901234567",\n'
        '  "password": "Abc@123456",\n'
        '  "otp_code": "888888"\n'
        '}'
    ),
    happy_response=(
        'HTTP 201 Created\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "access_token": "eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9...",\n'
        '    "refresh_token": "eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9...",\n'
        '    "user": {\n'
        '      "id": "550e8400-e29b-41d4-a716-446655440001",\n'
        '      "email": "nguyenvana@gmail.com",\n'
        '      "phone": "0901234567",\n'
        '      "full_name": "Nguyen Van A",\n'
        '      "role": "CUSTOMER",\n'
        '      "email_verified": false\n'
        '    }\n'
        '  }\n'
        '}\n'
        '// Side effects:\n'
        '// - Gửi email xác thực tới nguyenvana@gmail.com\n'
        '// - Ghi audit_log: action=USER_REGISTERED'
    ),
    error_responses=(
        'HTTP 409 Conflict (email/phone đã tồn tại)\n'
        '{ "success": false, "error": { "code": "ERR_REG_001", "message": "Email hoặc số điện thoại này đã được đăng ký" } }\n\n'
        'HTTP 400 Bad Request (OTP sai/hết hạn)\n'
        '{ "success": false, "error": { "code": "ERR_REG_002", "message": "Mã OTP không đúng hoặc đã hết hạn" } }\n\n'
        'HTTP 422 Unprocessable (password yếu)\n'
        '{ "success": false, "error": { "code": "ERR_REG_003", "message": "Mật khẩu không đủ mạnh", "details": [{"field":"password","msg":"Cần ít nhất 10 ký tự..."}] } }\n\n'
        'HTTP 429 Too Many Requests\n'
        '{ "success": false, "error": { "code": "ERR_RATE_LIMIT", "message": "Quá nhiều yêu cầu. Thử lại sau 1 giờ." } }'
    )
)

# ─── SCR-00-FP ───
add_h3(doc, '2.1.3. SCR-00-FP — Quên Mật Khẩu & Đặt Lại Mật Khẩu')
add_bullet(doc, 'URL: /auth/forgot-password  |  URL đặt lại: /auth/reset-password?token=xxx')
add_bullet(doc, 'Prerequisites: Đã có tài khoản. Flow: 2 bước (nhập identity → nhận OTP → đặt lại)')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [🚗 AUTO DEALERSHIP]                       QUÊN MẬT KHẨU               │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│                                                                          │\n'
'│  BƯỚC 1 — Xác minh danh tính                                            │\n'
'│                                                                          │\n'
'│  Email hoặc Số điện thoại *                                              │\n'
'│  [ identity_input                                         ]              │\n'
'│                                                                          │\n'
'│  [ GỬI MÃ XÁC NHẬN ] → hệ thống gửi OTP về SĐT hoặc link qua email    │\n'
'│                                                                          │\n'
'│  ─────────────────────────────────────────────                           │\n'
'│                                                                          │\n'
'│  BƯỚC 2 — Đặt lại mật khẩu (sau khi nhận OTP)                          │\n'
'│                                                                          │\n'
'│  Mã OTP *                                                                │\n'
'│  [ ][ ][ ][ ][ ][ ]   ⏱ Còn 300s (5 phút)                             │\n'
'│                                                                          │\n'
'│  Mật khẩu mới *                                       [👁]               │\n'
'│  [ new_password                                           ]              │\n'
'│                                                                          │\n'
'│  Xác nhận mật khẩu mới *                             [👁]               │\n'
'│  [ confirm_new_password                                   ]              │\n'
'│                                                                          │\n'
'│  [ ĐẶT LẠI MẬT KHẨU ]                                                  │\n'
'│                                                                          │\n'
'│  [ Quay lại đăng nhập ]                                                  │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-00-FP')
add_field_table(doc, [
    ('identity_input',    'Text Field',     'Bắt buộc', 'Email RFC5322 hoặc SĐT VN',
     'Nhận SĐT → gửi OTP SMS (TTL 5 phút). Nhận email → gửi link reset (TTL 15 phút). Không tiết lộ user có tồn tại hay không (security).'),
    ('btn_send_verify',   'Primary Button', 'N/A', 'N/A',
     'Disabled nếu identity_input rỗng hoặc sai format. Gọi POST /auth/forgot-password. Loading state khi đang gửi.'),
    ('otp_digits',        '6-box Input',    'Bắt buộc (flow OTP)', '6 chữ số. TTL 300s.',
     'Chỉ hiện sau khi bước 1 thành công. Auto-focus. Countdown 5 phút.'),
    ('new_password',      'Password Field', 'Bắt buộc', 'Cùng rule với password đăng ký: min 10, hoa, thường, số, đặc biệt.',
     'Không được trùng với mật khẩu cũ (backend check). Password strength indicator.'),
    ('confirm_new_pw',    'Password Field', 'Bắt buộc', 'Phải khớp new_password.',
     'Client-side validation, không gửi lên server.'),
    ('btn_reset',         'Primary Button', 'N/A', 'N/A',
     'Disabled nếu OTP chưa nhập đủ hoặc password không hợp lệ. Gọi POST /auth/reset-password. Sau thành công → redirect /auth/login.'),
    ('link_back_login',   'Text Link',      'N/A', 'N/A',
     'Navigate về /auth/login. Không cần confirm.'),
])

add_h3(doc, 'API Specs — SCR-00-FP')

add_api_spec(doc, 'POST', '/api/v1/auth/forgot-password',
    auth='Public. Rate limit: 3 req/identity/hour.',
    desc='Gửi OTP SMS hoặc link reset email. Không tiết lộ account có tồn tại hay không.',
    request_body='{\n  "identity": "0901234567"    // email hoặc phone\n}',
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "message": "Mã xác nhận đã được gửi",\n'
        '    "method": "SMS",        // hoặc "EMAIL"\n'
        '    "expires_in": 300,      // giây\n'
        '    "masked_target": "090*****67"\n'
        '  }\n'
        '}\n'
        '// Nếu identity không tồn tại: vẫn trả 200 (không leak info)'
    ),
    error_responses=(
        'HTTP 429 Too Many Requests\n'
        '{ "success": false, "error": { "code": "ERR_FP_RATE_LIMIT", "message": "Quá nhiều yêu cầu. Thử lại sau 1 giờ." } }\n\n'
        'HTTP 400 Bad Request\n'
        '{ "success": false, "error": { "code": "ERR_INVALID_IDENTITY", "message": "Email hoặc số điện thoại không hợp lệ" } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/auth/reset-password',
    auth='Public.',
    desc='Đặt lại mật khẩu. Token chỉ dùng được 1 lần, TTL 15 phút.',
    request_body=(
        '{\n'
        '  "identity": "0901234567",    // email hoặc phone\n'
        '  "otp_code": "123456",        // OTP từ SMS, hoặc\n'
        '  "token": "abc123xyz...",      // token từ link email\n'
        '  "new_password": "NewPass@99"\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "message": "Mật khẩu đã được đặt lại thành công. Vui lòng đăng nhập lại."\n'
        '  }\n'
        '}\n'
        '// Side effects:\n'
        '// - Invalidate tất cả user_sessions (force logout tất cả thiết bị)\n'
        '// - Ghi audit_log: action=PASSWORD_RESET'
    ),
    error_responses=(
        'HTTP 400 Bad Request (token/OTP hết hạn)\n'
        '{ "success": false, "error": { "code": "ERR_FP_TOKEN_EXPIRED", "message": "Mã xác nhận đã hết hạn. Vui lòng yêu cầu lại." } }\n\n'
        'HTTP 400 (token đã dùng)\n'
        '{ "success": false, "error": { "code": "ERR_FP_TOKEN_USED", "message": "Mã xác nhận đã được sử dụng." } }\n\n'
        'HTTP 422 (password trùng cũ)\n'
        '{ "success": false, "error": { "code": "ERR_SAME_PASSWORD", "message": "Mật khẩu mới không được trùng với mật khẩu cũ" } }'
    )
)

# ─── SCR-00-PROF ───
add_h3(doc, '2.1.4. SCR-00-PROF — Hồ Sơ Cá Nhân & Cập Nhật CCCD')
add_bullet(doc, 'URL: /profile  |  Auth: Bearer token. Role: All.')
add_bullet(doc, 'Prerequisites: Đã đăng nhập. customer_profile có thể chưa hoàn thiện.')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [🚗 AUTO DEALERSHIP]   Trang Chủ   [👤 Nguyen Van A ▾]  [Đăng Xuất]    │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│                          HỒ SƠ CÁ NHÂN                                  │\n'
'│                                                                          │\n'
'│  Thông tin cơ bản                    Xác thực danh tính                 │\n'
'│  ┌────────────────────────┐          ┌──────────────────────────────┐   │\n'
'│  │ Họ tên: [full_name  ] │          │ Số CCCD: [identity_card_num] │   │\n'
'│  │ Email:  [email      ] │          │          ••••••••••••         │   │\n'
'│  │ SĐT:    [phone      ] │          │ Ngày cấp: [datepicker      ] │   │\n'
'│  │ Địa chỉ:[address    ] │          │ Nơi cấp: [issue_place      ] │   │\n'
'│  │ Thu nhập:[monthly_  ] │          │                               │   │\n'
'│  │          income/tháng │          │ Upload ảnh CCCD:              │   │\n'
'│  │                       │          │ [+ Mặt trước]  [+ Mặt sau]   │   │\n'
'│  │ [ Lưu thông tin ]     │          │                               │   │\n'
'│  └────────────────────────┘          │ [ Cập nhật CCCD ]            │   │\n'
'│                                      └──────────────────────────────┘   │\n'
'│                                                                          │\n'
'│  Bảo mật tài khoản                                                       │\n'
'│  [ Đổi mật khẩu ]   [ Quản lý thiết bị đăng nhập ]                     │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-00-PROF')
add_field_table(doc, [
    ('full_name',              'Text Field',   'Bắt buộc', 'Min 2, Max 100 ký tự',
     'Tên hiển thị trên UI và hợp đồng. Cập nhật realtime trên header sau save.'),
    ('email',                  'Text Field',   'Bắt buộc', 'RFC5322. UNIQUE.',
     'Sau update email: gửi verification link. email_verified=false cho đến khi verify. Không cho đăng nhập bằng email mới cho đến khi verified.'),
    ('phone',                  'Text Field',   'Bắt buộc', '^(0|84)[35789][0-9]{8}$. UNIQUE.',
     'Cập nhật phone yêu cầu xác thực OTP về số mới.'),
    ('permanent_address',      'Textarea',     'Bắt buộc', 'Min 10, Max 255 ký tự',
     'Địa chỉ hộ khẩu thường trú ghi trên hợp đồng mua xe.'),
    ('monthly_income',         'Number Input', 'Tùy chọn', 'Min 0. Max 999,999,999. Đơn vị: VNĐ/tháng',
     'Dùng để hệ thống gợi ý giới hạn vay tự động (thường <= 40% thu nhập tháng). Hiển thị số có dấu phân cách ngàn.'),
    ('identity_card_number',   'Text Field',   'Bắt buộc', '12 chữ số. Structural: 3 số đầu = mã tỉnh 001-096',
     'AES-256-GCM encrypt trước khi lưu DB. UI Masking: 001xxx****5678. Chỉ Admin mới có thể decrypt xem full.'),
    ('identity_card_date',     'Datepicker',   'Bắt buộc', 'Date >= 2012-01-01 AND Date <= Today',
     'CCCD mới cấp từ năm 2012. Hiện calendar picker. Validate ngay khi blur.'),
    ('identity_card_place',    'Text Field',   'Bắt buộc', 'Min 5, Max 150 ký tự',
     'Nơi cấp CCCD (tên Cục/Phòng cảnh sát). Dùng trên hợp đồng.'),
    ('cccd_front_upload',      'File Input',   'Tùy chọn', 'JPG/PNG/HEIC/PDF. Max 5MB.',
     'Upload ảnh mặt trước CCCD. Lưu S3, URL vào customer_profiles.cccd_front_url. Preview thumbnail sau upload.'),
    ('cccd_back_upload',       'File Input',   'Tùy chọn', 'JPG/PNG/HEIC/PDF. Max 5MB.',
     'Upload ảnh mặt sau CCCD. Preview thumbnail sau upload.'),
    ('btn_save_basic',         'Primary Btn',  'N/A', 'N/A',
     'Gọi PUT /users/profile. Disabled nếu không có thay đổi. Loading state khi đang lưu. Toast success/error.'),
    ('btn_update_cccd',        'Primary Btn',  'N/A', 'N/A',
     'Gọi PUT /users/profile với CCCD fields. Ghi audit_log sau update. Hiện confirm dialog trước khi submit.'),
    ('btn_change_password',    'Ghost Button', 'N/A', 'N/A',
     'Mở modal đổi mật khẩu: nhập current_password + new_password + confirm. Gọi PUT /users/change-password.'),
    ('btn_manage_sessions',    'Ghost Button', 'N/A', 'N/A',
     'Hiện danh sách thiết bị đang đăng nhập (từ user_sessions). Cho phép logout từng thiết bị.'),
])

add_h3(doc, 'API Specs — SCR-00-PROF')
add_api_spec(doc, 'GET', '/api/v1/users/profile',
    auth='Bearer token. Role: All.',
    desc='Lấy thông tin hồ sơ của user đang đăng nhập.',
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "id": "550e8400-e29b-41d4-a716-446655440000",\n'
        '    "email": "user@example.com",\n'
        '    "email_verified": true,\n'
        '    "phone": "0901234567",\n'
        '    "full_name": "Nguyen Van A",\n'
        '    "role": "CUSTOMER",\n'
        '    "profile": {\n'
        '      "identity_card_masked": "001200001234",\n'
        '      "identity_card_date": "2020-05-15",\n'
        '      "identity_card_place": "Cuc Canh sat DKQL - CA TP HCM",\n'
        '      "permanent_address": "123 Nguyen Hue, Q1, TP.HCM",\n'
        '      "monthly_income": 25000000,\n'
        '      "cccd_front_url": "https://cdn.autodealer.vn/cccd/...",\n'
        '      "cccd_back_url": null\n'
        '    }\n'
        '  }\n'
        '}'
    ),
    error_responses='HTTP 401 Unauthorized\n{ "success": false, "error": { "code": "ERR_AUTH_REQUIRED", "message": "Vui lòng đăng nhập" } }'
)

add_api_spec(doc, 'PUT', '/api/v1/users/profile',
    auth='Bearer token. BOLA: chỉ update profile của chính mình.',
    desc='Cập nhật thông tin hồ sơ. Partial update (chỉ gửi fields cần đổi).',
    request_body=(
        '{\n'
        '  "full_name": "Nguyen Van B",             // optional\n'
        '  "permanent_address": "456 Le Loi, Q1",   // optional\n'
        '  "monthly_income": 30000000,              // optional\n'
        '  "identity_card_number": "001200001234",  // optional, triggers AES encrypt\n'
        '  "identity_card_date": "2020-05-15",      // optional\n'
        '  "identity_card_place": "Cuc Canh sat"   // optional\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "message": "Hồ sơ đã được cập nhật",\n'
        '    "updated_fields": ["full_name", "monthly_income"]\n'
        '  }\n'
        '}\n'
        '// Side effects: ghi audit_log với old_value và new_value (CCCD masked)'
    ),
    error_responses=(
        'HTTP 422 Unprocessable\n'
        '{ "success": false, "error": { "code": "ERR_VALIDATION", "message": "Dữ liệu không hợp lệ", "details": [{"field":"identity_card_number","msg":"Mã CCCD không đúng cấu trúc"}] } }\n\n'
        'HTTP 409 Conflict\n'
        '{ "success": false, "error": { "code": "ERR_PHONE_TAKEN", "message": "Số điện thoại đã được dùng bởi tài khoản khác" } }'
    )
)

add_api_spec(doc, 'PUT', '/api/v1/users/change-password',
    auth='Bearer token. BOLA: chỉ đổi mật khẩu của chính mình.',
    desc='Đổi mật khẩu. Yêu cầu xác nhận mật khẩu cũ.',
    request_body=(
        '{\n'
        '  "current_password": "OldPass@88",\n'
        '  "new_password": "NewPass@99"\n'
        '}'
    ),
    happy_response='HTTP 200 OK\n{\n  "success": true,\n  "data": { "message": "Mật khẩu đã được thay đổi. Vui lòng đăng nhập lại." }\n}\n// Side effect: invalidate tất cả user_sessions trừ session hiện tại',
    error_responses=(
        'HTTP 400 Bad Request\n{ "success": false, "error": { "code": "ERR_WRONG_PASSWORD", "message": "Mật khẩu hiện tại không đúng" } }\n\n'
        'HTTP 422 Unprocessable\n{ "success": false, "error": { "code": "ERR_SAME_PASSWORD", "message": "Mật khẩu mới không được trùng với mật khẩu cũ" } }'
    )
)

doc.add_page_break()

# ─── SCR-01 & SCR-02 ───
add_h2(doc, '2.2. Luồng 2: Danh Mục Xe & Chi Tiết Xe PDP (SCR-01 & SCR-02)')

add_h3(doc, '2.2.1. SCR-01 — Danh Mục Xe & Bộ Lọc')
add_bullet(doc, 'URL: /catalog  |  Auth: Public (không cần đăng nhập)')
add_bullet(doc, 'Navigation: Header nav → [Danh Mục Xe]. Chọn xe → SCR-02.')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [🚗 AUTO DEALERSHIP]  Trang Chủ  Danh Mục  Lái Thử  [🔍] [👤 Login]   │\n'
'├────────────┬────────────────────────────────────────────────────────────┤\n'
'│ BỘ LỌC     │  Kết quả: 24 xe   [Sắp xếp: Giá ▾]                       │\n'
'│            │                                                            │\n'
'│ Hãng xe    │  ┌──────────┐  ┌──────────┐  ┌──────────┐                │\n'
'│ ☐ Toyota  │  │ [img]    │  │ [img]    │  │ [img]    │                │\n'
'│ ☐ Honda   │  │ Toyota   │  │ Honda    │  │ Hyundai  │                │\n'
'│ ☐ Hyundai │  │ Camry 2.5│  │ CR-V     │  │ Santa Fe │                │\n'
'│            │  │ 1.29 tỷ  │  │ 1.05 tỷ  │  │ 1.35 tỷ  │                │\n'
'│ Phân khúc  │  │ [Xem chi]│  │ [Xem chi]│  │ [Xem chi]│                │\n'
'│ ○ Sedan   │  │  tiết]   │  │  tiết]   │  │  tiết]   │                │\n'
'│ ○ SUV     │  └──────────┘  └──────────┘  └──────────┘                │\n'
'│ ○ Pickup  │                                                            │\n'
'│            │  ← Trang 1 / 3 →                                         │\n'
'│ Khoảng giá │                                                            │\n'
'│ [──●──────]│  🎯 Quiz: Xe phù hợp với nhu cầu của bạn?                │\n'
'│ 500M - 2B  │  [🏠 Gia đình] [🏃 Thể thao] [💼 Kinh doanh] [🌱 Eco]    │\n'
'│            │                                                            │\n'
'│ Showroom   │                                                            │\n'
'│ [Chọn SR ▾]│                                                            │\n'
'└────────────┴────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-01')
add_field_table(doc, [
    ('filter_brand',         'Multi-Checkbox', 'Tùy chọn', 'Array of brand_id UUIDs',
     'Lọc theo hãng. URL sync: ?brand=toyota,honda. Kết quả refresh realtime (không cần bấm Apply). Reset khi chọn [Bỏ lọc].'),
    ('filter_segment',       'Radio Group',    'Tùy chọn', 'Enum: SEDAN|SUV|PICKUP|HATCHBACK|MPV|ELECTRIC',
     'Chỉ chọn 1 phân khúc. URL sync: ?segment=suv. Kết hợp với filter_brand.'),
    ('filter_price_range',   'Dual Slider',    'Tùy chọn', 'Min >= 0, Max <= 5,000,000,000 VNĐ',
     'Slider từ 500M đến 5B. Giá trị hiển thị format "X tỷ Y triệu". Debounce 300ms trước khi fetch.'),
    ('filter_showroom',      'Dropdown',       'Tùy chọn', 'showroom_id UUID',
     'Lọc xe available theo showroom. Fetch GET /showrooms để populate options. URL sync: ?showroom=uuid.'),
    ('sort_order',           'Dropdown',       'Tùy chọn', 'Enum: PRICE_ASC|PRICE_DESC|NEWEST|POPULAR',
     'Sắp xếp kết quả. Default: POPULAR. URL sync: ?sort=price_asc.'),
    ('search_keyword',       'Search Input',   'Tùy chọn', 'Max 100 chars',
     'Tìm theo tên model, phiên bản. Debounce 400ms. Highlight từ khóa trong kết quả.'),
    ('lifestyle_quiz',       'Button Group',   'Tùy chọn', 'N/A',
     '4 nút: Gia đình/Thể thao/Kinh doanh/Eco. Khi chọn → áp dụng preset filters. Gia đình→SUV/MPV. Thể thao→Sedan/Coupe. Eco→Electric/Hybrid.'),
    ('card_vehicle',         'Card Component', 'N/A', 'N/A',
     'Hiện: ảnh xe, tên model, giá từ, số lượng available. Hover → shadow elevation. Click → navigate SCR-02.'),
    ('btn_view_detail',      'Card Button',    'N/A', 'N/A',
     'Primary action trong card. Điều hướng → /catalog/{variant_id}. Có thể target="_blank" để mở tab mới.'),
    ('pagination',           'Pagination',     'N/A', 'N/A',
     'Hiện tối đa 12 xe/trang. Có prev/next và page numbers. URL sync: ?page=2.'),
])

add_h3(doc, 'API Specs — SCR-01')
add_api_spec(doc, 'GET', '/api/v1/catalog/models',
    auth='Public.',
    desc='Lấy danh sách models với filters và pagination.',
    query_params=(
        '?brand=uuid1,uuid2     // filter by brand_id (multiple)\n'
        '&segment=SUV           // filter by segment\n'
        '&min_price=500000000   // min price in VND\n'
        '&max_price=2000000000  // max price in VND\n'
        '&showroom_id=uuid      // filter by availability at showroom\n'
        '&sort=PRICE_ASC        // PRICE_ASC|PRICE_DESC|NEWEST|POPULAR\n'
        '&keyword=camry         // search keyword\n'
        '&page=1&limit=12       // pagination'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": [\n'
        '    {\n'
        '      "id": "model-uuid-001",\n'
        '      "brand": { "id": "brand-uuid", "name": "Toyota", "logo_url": "https://..." },\n'
        '      "name": "Camry",\n'
        '      "segment": "SEDAN",\n'
        '      "variants": [\n'
        '        {\n'
        '          "id": "variant-uuid-001",\n'
        '          "name": "Camry 2.5Q",\n'
        '          "listed_price": 1290000000,\n'
        '          "thumbnail_url": "https://cdn.autodealer.vn/...",\n'
        '          "available_count": 3,\n'
        '          "colors": ["Trắng Ngọc Trai", "Đen Ánh Kim"]\n'
        '        }\n'
        '      ]\n'
        '    }\n'
        '  ],\n'
        '  "meta": { "page": 1, "limit": 12, "total": 24, "total_pages": 2 }\n'
        '}'
    ),
    error_responses='HTTP 400 Bad Request\n{ "success": false, "error": { "code": "ERR_INVALID_FILTER", "message": "Bộ lọc không hợp lệ", "details": [{"field":"segment","msg":"Phân khúc không hợp lệ"}] } }'
)

add_api_spec(doc, 'GET', '/api/v1/showrooms',
    auth='Public.',
    desc='Lấy danh sách tất cả showrooms đang hoạt động.',
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": [\n'
        '    {\n'
        '      "id": "showroom-uuid-001",\n'
        '      "name": "AutoDealership Quận 1",\n'
        '      "code": "SR-Q1",\n'
        '      "address": "123 Nguyen Hue, Q1, TP.HCM",\n'
        '      "phone": "028-1234-5678",\n'
        '      "operating_hours": {\n'
        '        "Mon-Fri": "08:00-20:00",\n'
        '        "Sat-Sun": "08:00-18:00"\n'
        '      }\n'
        '    }\n'
        '  ]\n'
        '}'
    ),
    error_responses='HTTP 500 Internal Server Error\n{ "success": false, "error": { "code": "ERR_SERVER", "message": "Lỗi hệ thống. Vui lòng thử lại." } }'
)

# SCR-02
add_h3(doc, '2.2.2. SCR-02 — Chi Tiết Xe PDP (Product Detail Page)')
add_bullet(doc, 'URL: /catalog/{variant_id}  |  Auth: Public. Một số hành động (Đặt cọc, Lái thử) yêu cầu login.')
add_bullet(doc, 'Navigation: SCR-01 → SCR-02. Nút [Đặt cọc] → SCR-04. [Đặt lịch lái thử] → SCR-03. [Thu cũ đổi mới] → mở form SCR-02-TI.')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [🚗 AUTO DEALERSHIP]   Trang Chủ > Danh Mục > Toyota Camry 2.5Q        │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│  ┌──────────────────────────┐   ┌─────────────────────────────────────┐│\n'
'│  │  [360° Viewer]           │   │ Toyota Camry 2.5Q                   ││\n'
'│  │  🔄 Xoay 360°           │   │ ⭐ 4.8 (128 đánh giá)               ││\n'
'│  │  [img xe]               │   │                                     ││\n'
'│  │  ← [màu: Trắng] [Đen] →│   │ Giá niêm yết: 1,290,000,000 đ       ││\n'
'│  └──────────────────────────┘   │ Đặt cọc tối thiểu: 50,000,000 đ   ││\n'
'│                                 │                                     ││\n'
'│  Thông số kỹ thuật              │ Showroom: [Chọn showroom ▾]         ││\n'
'│  ┌──────────────────────────┐   │ Còn lại: 3 xe                       ││\n'
'│  │ Động cơ: 2.5L 4 xylanh │   │                                     ││\n'
'│  │ Công suất: 209 HP       │   │ [ ĐẶT CỌC NGAY ] 🔵                 ││\n'
'│  │ Hộp số: CVT             │   │ [ ĐẶT LỊCH LÁI THỬ ] 🔲             ││\n'
'│  │ Tiêu hao: 7.2L/100km   │   │ [ THU CŨ ĐỔI MỚI ] 🔲               ││\n'
'│  └──────────────────────────┘   │                                     ││\n'
'│                                 │ 🔒 Thanh toán bảo mật               ││\n'
'│  [Tính toán trả góp]            └─────────────────────────────────────┘│\n'
'│  Vay: [70%  ▾]  Kỳ hạn: [48 tháng ▾]  → Góp: ~15.2M/tháng           │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-02')
add_field_table(doc, [
    ('viewer_360',        '360° Image Viewer', 'N/A', 'N/A',
     'Cho phép xoay xe 360°. 36 frames ảnh. Swipe gesture trên mobile. Nút [🔄 Xoay 360°] kích hoạt auto-rotate.'),
    ('color_selector',   'Image Swatch',      'Tùy chọn', 'color_id string',
     'Chọn màu xe → cập nhật ảnh viewer. URL sync: ?color=trang-ngoc-trai. Hiện số lượng available theo màu.'),
    ('showroom_selector', 'Dropdown',          'Bắt buộc trước đặt cọc', 'showroom_id UUID',
     'Chọn showroom nhận xe. Sau chọn → fetch quota realtime. Nếu hết xe → disable [Đặt cọc]. Lưu vào localStorage.'),
    ('available_count',   'Live Counter',      'N/A', 'N/A',
     'Hiển thị số xe còn lại tại showroom đã chọn. Fetch từ API. Cập nhật mỗi 30s (polling). Màu đỏ nếu <= 1.'),
    ('spec_table',        'Data Table',        'N/A', 'N/A',
     'Hiển thị specs_json. Expand/collapse các nhóm: Động cơ, Khung gầm, Tiện nghi, An toàn. Responsive: accordion trên mobile.'),
    ('btn_deposit',       'Primary Button',    'N/A', 'N/A',
     'Yêu cầu đăng nhập (CUSTOMER). Yêu cầu chọn showroom. Disabled nếu available_count=0. Click → /checkout/{variant_id}.'),
    ('btn_test_drive',    'Secondary Button',  'N/A', 'N/A',
     'Yêu cầu đăng nhập. Click → mở modal hoặc /test-drive với variant_id preset. Available khi showroom có slot trống.'),
    ('btn_trade_in',      'Secondary Button',  'N/A', 'N/A',
     'Mở form thu cũ đổi mới (SCR-02-TI) như slide-in panel. User nhập thông tin xe cũ → nhận định giá sơ bộ.'),
    ('loan_calculator',   'Interactive Widget','N/A', 'N/A',
     'Slider vay (10%-90%). Dropdown kỳ hạn (12-96 tháng). Real-time tính: monthly = principal * rate / (1-(1+rate)^-n). Hiện lãi suất hiện tại.'),
])

add_h3(doc, 'API Specs — SCR-02')
add_api_spec(doc, 'GET', '/api/v1/catalog/variants/{variant_id}',
    auth='Public.',
    desc='Lấy chi tiết một phiên bản xe theo ID.',
    path_params='variant_id: UUID — ID của vehicle_variant',
    query_params='?color=color-code  // optional, lọc ảnh theo màu\n&showroom_id=uuid // optional, lấy quota tại showroom',
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "id": "variant-uuid-001",\n'
        '    "model": { "id": "model-uuid", "name": "Camry", "brand": "Toyota" },\n'
        '    "name": "Camry 2.5Q",\n'
        '    "listed_price": 1290000000,\n'
        '    "min_deposit_amount": 50000000,\n'
        '    "specs_json": {\n'
        '      "engine": "2.5L 4-cylinder",\n'
        '      "power_hp": 209,\n'
        '      "transmission": "CVT",\n'
        '      "fuel_consumption": "7.2L/100km"\n'
        '    },\n'
        '    "images": [\n'
        '      { "color": "Trắng Ngọc Trai", "urls": ["https://cdn.../1.jpg","https://cdn.../2.jpg"] }\n'
        '    ],\n'
        '    "quota": {\n'
        '      "showroom_id": "showroom-uuid",\n'
        '      "available_count": 3,\n'
        '      "soft_locked_count": 1\n'
        '    }\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 404 Not Found\n{ "success": false, "error": { "code": "ERR_VARIANT_NOT_FOUND", "message": "Không tìm thấy xe này" } }'
    )
)

doc.add_page_break()

# SCR-03
add_h2(doc, '2.3. Luồng 3: Đặt Lịch Lái Thử (SCR-03)')
add_bullet(doc, 'URL: /test-drive  |  Auth: CUSTOMER (đăng nhập). SALE có thể đặt hộ.')
add_bullet(doc, 'Prerequisites: Đã đăng nhập. Showroom có slots trống. customer_profile đã điền CCCD.')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [🚗 AUTO DEALERSHIP]                ĐĂNG KÝ TRẢI NGHIỆM LÁI THỬ XE    │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│                                                                          │\n'
'│  Xe muốn lái thử: Toyota Camry 2.5Q  [Thay đổi]                        │\n'
'│                                                                          │\n'
'│  Showroom *                                                              │\n'
'│  [ Chọn showroom                                        ▾ ]             │\n'
'│                                                                          │\n'
'│  Ngày lái thử *                                                          │\n'
'│  [ 📅 Thứ Ba, 25/08/2026                                ▾ ]             │\n'
'│                                                                          │\n'
'│  Khung giờ *                                                             │\n'
'│  [ 09:00-10:00 ✅ ]  [ 10:00-11:00 ❌ ]  [ 14:00-15:00 ✅ ]             │\n'
'│  [ 15:00-16:00 ✅ ]  [ 16:00-17:00 ❌ ]  [ 17:00-18:00 ✅ ]             │\n'
'│                                                                          │\n'
'│  Thông tin người lái:                                                    │\n'
'│  ○ Chính tôi lái    ○ Người thân lái (nhập thêm thông tin)              │\n'
'│                                                                          │\n'
'│  Số GPLX *                                                               │\n'
'│  [ driver_license                                         ]             │\n'
'│  ⓘ Giấy phép lái xe 12 chữ số                                          │\n'
'│                                                                          │\n'
'│  Ghi chú                                                                 │\n'
'│  [ notes                                                  ]             │\n'
'│                                                                          │\n'
'│  [           XÁC NHẬN ĐẶT LỊCH           ] 🔵                          │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-03')
add_field_table(doc, [
    ('variant_display',    'Read-only',      'N/A', 'N/A',
     'Hiện tên xe đang muốn lái thử. Lấy từ URL param hoặc session. Nút [Thay đổi] → /catalog để chọn lại.'),
    ('showroom_selector',  'Dropdown',       'Bắt buộc', 'showroom_id UUID',
     'Sau chọn → fetch slots available cho showroom đó theo ngày đã chọn. Hiện showroom gần vị trí user nếu có geolocation.'),
    ('date_picker',        'Datepicker',     'Bắt buộc', 'Date >= Today + 1 ngày. Chỉ ngày showroom mở cửa.',
     'Không cho chọn ngày đã qua hoặc ngày showroom đóng cửa (dựa showroom_operating_hours). Fetch slots khi chọn ngày.'),
    ('slot_selector',      'Button Grid',    'Bắt buộc', 'slot_id UUID',
     '✅ Slot available (xanh). ❌ Slot đã đặt (xám, disabled). Chỉ chọn 1 slot. Slot hết chỗ sau 5 phút không confirm (optimistic lock).'),
    ('driver_type',        'Radio Group',    'Bắt buộc', 'Enum: SELF | BEHALF',
     'SELF: dùng thông tin user đang login. BEHALF: hiện thêm fields nhập họ tên, SĐT, GPLX của người thân.'),
    ('driver_license',     'Text Field',     'Bắt buộc', '^[0-9]{12}$. 12 chữ số.',
     'Số Giấy Phép Lái Xe bắt buộc. Validate format 12 chữ số. Lưu vào test_drive_bookings.driver_license.'),
    ('behalf_name',        'Text Field',     'Bắt buộc nếu BEHALF', 'Min 2, Max 100 ký tự',
     'Chỉ hiện khi driver_type=BEHALF. Tên người sẽ lái xe thực tế.'),
    ('behalf_phone',       'Text Field',     'Bắt buộc nếu BEHALF', '^(0|84)[35789][0-9]{8}$',
     'SĐT người lái hộ để Sale liên hệ nếu cần.'),
    ('notes',              'Textarea',       'Tùy chọn', 'Max 500 ký tự',
     'Ghi chú thêm (yêu cầu đặc biệt, điểm đón...). Hiện counter ký tự.'),
    ('btn_confirm',        'Primary Button', 'N/A', 'N/A',
     'Disabled nếu thiếu required fields. Gọi POST /test-drives/book. Loading state. Sau thành công → hiện confirmation modal + gửi SMS.'),
])

add_h3(doc, 'API Specs — SCR-03')
add_api_spec(doc, 'GET', '/api/v1/test-drives/slots',
    auth='Public.',
    desc='Lấy danh sách slots lái thử available theo showroom và ngày.',
    query_params=(
        '?showroom_id=uuid    // required\n'
        '&date=2026-08-25     // required, format YYYY-MM-DD\n'
        '&variant_id=uuid     // optional, filter by xe'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": [\n'
        '    {\n'
        '      "id": "slot-uuid-001",\n'
        '      "slot_start": "2026-08-25T09:00:00+07:00",\n'
        '      "slot_end": "2026-08-25T10:00:00+07:00",\n'
        '      "is_booked": false,\n'
        '      "demo_vehicle": { "vin": "VN1234567890", "color": "Trắng" },\n'
        '      "assigned_sale": { "id": "sale-uuid", "name": "Tran Van B" }\n'
        '    },\n'
        '    {\n'
        '      "id": "slot-uuid-002",\n'
        '      "slot_start": "2026-08-25T10:00:00+07:00",\n'
        '      "slot_end": "2026-08-25T11:00:00+07:00",\n'
        '      "is_booked": true,\n'
        '      "demo_vehicle": null,\n'
        '      "assigned_sale": null\n'
        '    }\n'
        '  ]\n'
        '}'
    ),
    error_responses=(
        'HTTP 400 Bad Request\n{ "success": false, "error": { "code": "ERR_INVALID_DATE", "message": "Không thể lấy slots cho ngày trong quá khứ" } }\n\n'
        'HTTP 404 Not Found\n{ "success": false, "error": { "code": "ERR_SHOWROOM_NOT_FOUND", "message": "Showroom không tồn tại" } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/test-drives/book',
    auth='Bearer token. Role: CUSTOMER | SALE.',
    desc='Đặt lịch lái thử. Tạo bản ghi test_drive_bookings. Lock slot.',
    request_body=(
        '{\n'
        '  "slot_id": "slot-uuid-001",\n'
        '  "variant_id": "variant-uuid-001",\n'
        '  "driver_license": "123456789012",   // required, 12 digits\n'
        '  "driver_type": "SELF",              // SELF | BEHALF\n'
        '  // Nếu driver_type = BEHALF:\n'
        '  "on_behalf_customer_name": "Le Thi C",\n'
        '  "on_behalf_customer_phone": "0907654321",\n'
        '  "notes": "Muốn test đường cao tốc"\n'
        '}'
    ),
    happy_response=(
        'HTTP 201 Created\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "booking_id": "booking-uuid-001",\n'
        '    "status": "CONFIRMED",\n'
        '    "slot": {\n'
        '      "date": "2026-08-25",\n'
        '      "time": "09:00-10:00",\n'
        '      "showroom": "AutoDealership Quận 1",\n'
        '      "sale_name": "Tran Van B",\n'
        '      "sale_phone": "0901111222"\n'
        '    },\n'
        '    "vehicle": { "name": "Toyota Camry 2.5Q", "color": "Trắng Ngọc Trai" }\n'
        '  }\n'
        '}\n'
        '// Side effects:\n'
        '// - Gửi SMS xác nhận tới customer\n'
        '// - Gửi SMS reminder trước 2 giờ (scheduled job)\n'
        '// - Mark slot is_booked = true'
    ),
    error_responses=(
        'HTTP 409 Conflict (slot đã bị book)\n'
        '{ "success": false, "error": { "code": "ERR_SLOT_TAKEN", "message": "Khung giờ này đã có người đặt. Vui lòng chọn khung giờ khác." } }\n\n'
        'HTTP 400 Bad Request (GPLX không hợp lệ)\n'
        '{ "success": false, "error": { "code": "ERR_INVALID_GPLX", "message": "Số GPLX phải là 12 chữ số" } }\n\n'
        'HTTP 403 Forbidden\n'
        '{ "success": false, "error": { "code": "ERR_PROFILE_INCOMPLETE", "message": "Vui lòng hoàn thiện hồ sơ cá nhân trước khi đặt lịch lái thử" } }'
    )
)

doc.add_page_break()

# SCR-04 & SCR-10
add_h2(doc, '2.4. Luồng 4: Đặt Cọc & Thanh Toán Mock Sandbox (SCR-04 → SCR-10)')
add_bullet(doc, 'URL: /checkout/{variant_id}  |  Auth: CUSTOMER. SALE có thể tạo hộ.')
add_bullet(doc, 'Prerequisites: Đăng nhập. Hồ sơ hoàn thiện (CCCD). Chọn showroom. available_count > 0.')

add_wireframe(doc,
'┌──────────────────────────────┬──────────────────────────────────────────┐\n'
'│  THÔNG TIN ĐƠN HÀNG         │  THANH TOÁN ĐẶT CỌC AN TOÀN             │\n'
'├──────────────────────────────┤                                          │\n'
'│ 🚗 Toyota Camry 2.5Q        │  Phương thức thanh toán:                 │\n'
'│ Màu: Trắng Ngọc Trai        │  ○ QR VietQR  ○ VNPay  ○ Chuyển khoản  │\n'
'│ Showroom: Q1                │                                          │\n'
'│                              │  ┌──────────────────────────────────┐   │\n'
'│ Giá niêm yết: 1,290,000,000 │  │  [QR CODE dynamically generated] │   │\n'
'│ Phụ kiện:    +  5,000,000  │  │                                   │   │\n'
'│ Bảo hiểm:   + 12,000,000   │  │  Ngân hàng: VCB                   │   │\n'
'│ Xe cũ đổi:  - 200,000,000  │  │  STK: 1234 5678 9012              │   │\n'
'│ Ví tín dụng: - 10,000,000  │  │  Nội dung: ORD-2026-001           │   │\n'
'│ ─────────────────────────── │  │  Số tiền: 50,000,000 đ            │   │\n'
'│ Tổng cọc:    50,000,000     │  └──────────────────────────────────┘   │\n'
'│                              │                                          │\n'
'│ Phụ kiện:                   │  ⏱ Còn lại: 14:32 để hoàn tất          │\n'
'│ [+ Thêm phụ kiện]           │                                          │\n'
'│                              │  📡 Đang chờ xác nhận thanh toán...     │\n'
'│ Bảo hiểm:                   │  (SSE realtime listener)                 │\n'
'│ [○ Bảo hiểm toàn diện 12M]  │                                          │\n'
'│ [○ Không cần bảo hiểm    ]  │  Hoặc: [Quay lại chọn phương thức]     │\n'
'│                              │         [Hủy đặt cọc]                   │\n'
'│ Ví tín dụng:                │                                          │\n'
'│ Số dư: 10,000,000đ          │                                          │\n'
'│ [✅ Áp dụng ví tín dụng]    │                                          │\n'
'└──────────────────────────────┴──────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-04')
add_field_table(doc, [
    ('order_summary',          'Read-only Panel', 'N/A', 'N/A',
     'Hiển thị tóm tắt đơn hàng: xe, màu, showroom, giá. Tính toán realtime khi thêm accessories/insurance/trade-in/credit.'),
    ('accessories_picker',     'Multi-Select',    'Tùy chọn', 'Array of accessory_id UUID',
     'Lấy danh sách từ GET /accessories. Mỗi item hiện: tên, giá, ảnh. Cộng vào order_total ngay khi chọn. Cho phép chọn nhiều.'),
    ('insurance_picker',       'Radio Group',     'Tùy chọn', 'insurance_product_id UUID hoặc null',
     'Options lấy từ GET /insurance-products. Hiện: tên, hãng bảo hiểm, phí/năm, coverage type. Chỉ chọn 1 hoặc không chọn.'),
    ('credit_wallet',          'Toggle + Display','Tùy chọn', 'credit_amount_to_apply >= 0',
     'Chỉ hiện nếu customer có credit_balance > 0. Toggle áp dụng/không áp dụng. Số credit trừ vào deposit_amount realtime.'),
    ('deposit_amount',         'Number Display',  'N/A', 'N/A',
     'Số tiền cọc = max(min_deposit_amount, deposit_amount_from_config). Không cho nhập tự do (trừ khi config cho phép linh hoạt). Hiển thị nổi bật.'),
    ('payment_method',         'Radio Group',     'Bắt buộc', 'Enum: VIETQR|VNPAY|TRANSFER',
     'Chọn phương thức thanh toán. Khi chọn → generate QR code hoặc chuyển redirect. Lưu preference.'),
    ('qr_display',             'QR Code Widget',  'N/A', 'N/A',
     'Dynamic QR theo chuẩn VietQR. Tự refresh sau 5 phút. Nút [Tải ảnh QR] để save về máy. Hiện đầy đủ: ngân hàng, STK, nội dung, số tiền.'),
    ('countdown_timer',        'Live Countdown',  'N/A', 'N/A',
     'Đếm ngược từ 15 phút. Khi hết → order CANCELED tự động (scheduler). Màu đỏ khi < 3 phút. Realtime sync với server.'),
    ('sse_status_listener',    'SSE Component',   'N/A', 'N/A',
     'Kết nối GET /notifications/stream. Listen event payment_status_changed. Khi nhận → cập nhật UI ngay, không cần reload. Tự reconnect nếu mất kết nối.'),
    ('btn_cancel_deposit',     'Danger Button',   'N/A', 'N/A',
     'Hủy đơn → confirm dialog "Bạn có chắc muốn hủy?". Gọi DELETE /orders/{id}. Hoàn credit nếu đã áp dụng.'),
])

add_h3(doc, 'API Specs — SCR-04')
add_api_spec(doc, 'POST', '/api/v1/orders/deposit',
    auth='Bearer token. Role: CUSTOMER | SALE.',
    desc='Tạo đơn hàng đặt cọc. Lock quota (atomic SQL). Tạo payment record. Trả về thông tin thanh toán.',
    request_body=(
        '{\n'
        '  "variant_id": "variant-uuid-001",\n'
        '  "showroom_id": "showroom-uuid-001",\n'
        '  "purchase_type": "DIRECT",          // DIRECT | AUTO_LOAN\n'
        '  "deposit_amount": 50000000,         // VNĐ, >= min_deposit_amount\n'
        '  "payment_method": "VIETQR",         // VIETQR | VNPAY | TRANSFER\n'
        '  "accessories": [\n'
        '    { "accessory_id": "acc-uuid-001", "quantity": 1 }\n'
        '  ],\n'
        '  "insurance_product_id": "ins-uuid-001",  // optional\n'
        '  "trade_in_request_id": "ti-uuid-001",     // optional\n'
        '  "credit_amount_to_apply": 10000000        // optional, <= credit_balance\n'
        '}'
    ),
    happy_response=(
        'HTTP 201 Created\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "order_id": "order-uuid-001",\n'
        '    "order_code": "ORD-20260823-0001",\n'
        '    "status": "PENDING_PAYMENT",\n'
        '    "expires_at": "2026-08-23T02:45:00+07:00",\n'
        '    "payment": {\n'
        '      "transaction_ref": "TXN-20260823-001",\n'
        '      "method": "VIETQR",\n'
        '      "amount": 50000000,\n'
        '      "qr_data": "00020101...",\n'
        '      "bank_account": "1234 5678 9012",\n'
        '      "bank_name": "Vietcombank",\n'
        '      "transfer_note": "ORD-20260823-0001"\n'
        '    },\n'
        '    "order_summary": {\n'
        '      "listed_price": 1290000000,\n'
        '      "accessories_total": 5000000,\n'
        '      "insurance_total": 12000000,\n'
        '      "trade_in_credit": 200000000,\n'
        '      "credit_applied": 10000000,\n'
        '      "final_price": 1097000000\n'
        '    }\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 409 Conflict (hết xe - quota=0)\n'
        '{ "success": false, "error": { "code": "ERR_QUOTA_EXCEEDED", "message": "Xe này đã hết tại showroom bạn chọn. Vui lòng chọn showroom khác hoặc đặt lịch nhận xe." } }\n\n'
        'HTTP 400 Bad Request (deposit không hợp lệ)\n'
        '{ "success": false, "error": { "code": "ERR_DEPOSIT_TOO_LOW", "message": "Số tiền đặt cọc tối thiểu là 50,000,000 đ" } }\n\n'
        'HTTP 403 Forbidden (hồ sơ chưa hoàn thiện)\n'
        '{ "success": false, "error": { "code": "ERR_PROFILE_INCOMPLETE", "message": "Vui lòng hoàn thiện hồ sơ CCCD trước khi đặt cọc" } }\n\n'
        'HTTP 429 Too Many Requests\n'
        '{ "success": false, "error": { "code": "ERR_CHECKOUT_RATE_LIMIT", "message": "Bạn đã tạo quá nhiều đơn hôm nay. Vui lòng thử lại sau." } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/payments/mock-webhook',
    auth='Internal. HMAC-SHA256 signature verification (header: X-Webhook-Signature).',
    desc='Mock webhook từ payment gateway. Update payment record và order status theo State Machine.',
    request_body=(
        '{\n'
        '  "transaction_ref": "TXN-20260823-001",\n'
        '  "gateway": "MOCK_GATEWAY",\n'
        '  "status": "SUCCESS",              // SUCCESS | FAILED | PARTIAL\n'
        '  "received_amount": 50000000,\n'
        '  "gateway_transaction_no": "VCB123456789",\n'
        '  "gateway_response_code": "00",\n'
        '  "gateway_bank_code": "VCB",\n'
        '  "timestamp": "2026-08-23T02:31:00+07:00"\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "order_id": "order-uuid-001",\n'
        '    "order_status": "DEPOSIT_PAID",\n'
        '    "payment_status": "SUCCESS",\n'
        '    "sse_event_sent": true\n'
        '  }\n'
        '}\n'
        '// Side effects:\n'
        '// - Order status: PENDING_PAYMENT → DEPOSIT_PAID\n'
        '// - Ghi order_status_history\n'
        '// - Push SSE event tới client\n'
        '// - Ghi outbox_event: SEND_PAYMENT_SUCCESS_NOTIFICATION'
    ),
    error_responses=(
        'HTTP 401 Unauthorized (signature sai)\n'
        '{ "success": false, "error": { "code": "ERR_WEBHOOK_SIGNATURE", "message": "Webhook signature không hợp lệ" } }\n\n'
        'HTTP 404 Not Found\n'
        '{ "success": false, "error": { "code": "ERR_TRANSACTION_NOT_FOUND", "message": "Không tìm thấy giao dịch" } }\n\n'
        'HTTP 409 Conflict (idempotency)\n'
        '{ "success": false, "error": { "code": "ERR_DUPLICATE_WEBHOOK", "message": "Webhook này đã được xử lý" } }'
    )
)

doc.add_page_break()

# SCR-05 Dashboard
add_h2(doc, '2.5. Dashboard KPI (SCR-05)')
add_bullet(doc, 'URL: /portal/dashboard  |  Auth: ADMIN, MANAGER, SALE')
add_bullet(doc, 'Mỗi role thấy data phù hợp: Admin → toàn hệ thống, Manager → Showroom, Sale → cá nhân')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [DEALER PORTAL]  Dashboard  Kho Xe  Duyệt Giá  CRM  [👤 Admin]        │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│  Dashboard KPI — Tháng 08/2026  [Lọc: Showroom ▾] [Xuất báo cáo 📥]  │\n'
'│                                                                          │\n'
'│  ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────┐  │\n'
'│  │ 💰 Doanh thu │ │ 📈 Gross Mgn │ │ 🚗 Xe bán    │ │ 📋 Đơn hàng  │  │\n'
'│  │  4.2 tỷ      │ │   18.5%      │ │    12 xe     │ │    47 đơn    │  │\n'
'│  │ +12% vs T7   │ │ +2.1pp T8   │ │ +3 vs T7     │ │ 8 chờ duyệt  │  │\n'
'│  └──────────────┘ └──────────────┘ └──────────────┘ └──────────────┘  │\n'
'│                                                                          │\n'
'│  ┌──────────────┐ ┌──────────────┐                                      │\n'
'│  │ 🏆 Conv Rate │ │ ⏱ SLA Leads │                                      │\n'
'│  │   24.5%      │ │  6.2 giờ    │                                      │\n'
'│  │ Lead→ Order  │ │ avg response │                                      │\n'
'│  └──────────────┘ └──────────────┘                                      │\n'
'│                                                                          │\n'
'│  Top xe bán chạy tháng này:          Trạng thái đơn hàng:              │\n'
'│  1. Toyota Camry 2.5Q (5 xe)         ● PENDING: 8   ● PROCESSING: 12  │\n'
'│  2. Honda CR-V 1.5L (4 xe)           ● DELIVERED: 7  ● COMPLETED: 20  │\n'
'│  3. Hyundai Santa Fe (3 xe)           ● CANCELED: 5                    │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-05')
add_field_table(doc, [
    ('filter_month',        'Month Picker',    'Tùy chọn', 'YYYY-MM format',
     'Default: tháng hiện tại. Khi đổi → refetch tất cả KPI cards và charts. URL sync: ?month=2026-08.'),
    ('filter_showroom',     'Dropdown',        'Tùy chọn', 'showroom_id UUID. Admin only.',
     'Chỉ Admin thấy dropdown này. Manager chỉ thấy showroom của mình. Sale chỉ thấy data cá nhân.'),
    ('kpi_revenue',         'KPI Card',        'N/A', 'N/A',
     'Tổng doanh thu (SUM final_price của orders COMPLETED trong kỳ). So sánh % với tháng trước. Màu xanh = tăng, đỏ = giảm.'),
    ('kpi_gross_margin',    'KPI Card',        'N/A', 'N/A',
     'Gross margin % = (Revenue - COGS) / Revenue. Chỉ Admin/Manager thấy. Sale thấy hoa hồng cá nhân.'),
    ('kpi_cars_sold',       'KPI Card',        'N/A', 'N/A',
     'Số xe đã bán (orders COMPLETED). Admin: toàn hệ. Manager: showroom. Sale: cá nhân đã chốt.'),
    ('kpi_orders',          'KPI Card',        'N/A', 'N/A',
     'Tổng số đơn hàng trong kỳ. Badge phụ: số đơn đang chờ xử lý (PENDING_PAYMENT, BANK_APPROVING...).'),
    ('kpi_conversion_rate', 'KPI Card',        'N/A', 'N/A',
     'Tỷ lệ Lead → Order thành công. Công thức: orders_created / leads_assigned * 100%.'),
    ('kpi_sla_leads',       'KPI Card',        'N/A', 'N/A',
     'Thời gian phản hồi lead trung bình (giờ). Màu đỏ nếu > 24h (vi phạm SLA).'),
    ('chart_top_cars',      'List/Chart',      'N/A', 'N/A',
     'Top 5 xe bán chạy nhất trong kỳ. Có thể click vào xe để xem chi tiết đơn hàng.'),
    ('chart_order_status',  'Donut Chart',     'N/A', 'N/A',
     'Phân bổ đơn hàng theo trạng thái. Click vào segment → filter danh sách đơn theo status đó.'),
    ('btn_export',          'Secondary Button','N/A', 'N/A',
     'Xuất báo cáo Excel/PDF. Gọi GET /reports/dashboard?format=xlsx&month=2026-08. Download trực tiếp.'),
])

doc.add_page_break()

# SCR-06 CRM
add_h2(doc, '2.6. Luồng 9: CRM Leads Kanban (SCR-06)')
add_bullet(doc, 'URL: /portal/crm  |  Auth: SALE, MANAGER, ADMIN')
add_bullet(doc, 'Drag-and-drop Kanban. Quick-action bar. SLA countdown per lead.')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [DEALER PORTAL]  ...  [LEADS CRM]  ...  [👤 Sale: Tran Van B]         │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│  CRM LEADS   [🔍 Tìm kiếm lead...]  [Score ▾]  [+ Thêm Lead]          │\n'
'│                                                                          │\n'
'│  ┌──NEW (12)────┐ ┌──CONTACTED─┐ ┌──TEST DRIVE┐ ┌──NEGOTIATING┐       │\n'
'│  │ ┌──────────┐ │ │ ┌────────┐ │ │ ┌────────┐ │ │ ┌──────────┐ │      │\n'
'│  │ │Nguyen VA │ │ │ │Le TB   │ │ │ │Tran TC │ │ │ │Pham VD   │ │      │\n'
'│  │ │Camry 2.5Q│ │ │ │CR-V 1.5│ │ │ │Santa Fe│ │ │ │Tucson 2.0│ │      │\n'
'│  │ │⏱ 2h23m   │ │ │ │⏱ 1d2h  │ │ │ │✅ 25/8 │ │ │ │Giá 1.15B │ │      │\n'
'│  │ │Score: 85 │ │ │ │Score:72│ │ │ │Score:91│ │ │ │Score: 68 │ │      │\n'
'│  │ │[📞][💳]   │ │ │ │[📞][💳]│ │ │ │[📞][💳]│ │ │ │[📞][💳]  │ │      │\n'
'│  │ └──────────┘ │ │ └────────┘ │ │ └────────┘ │ │ └──────────┘ │      │\n'
'│  │ ┌──────────┐ │ │            │ │            │ │              │      │\n'
'│  │ │Hoang VE  │ │ │            │ │            │ │              │      │\n'
'│  │ │Accent 1.4│ │ │            │ │            │ └────────────── ┘     │\n'
'│  │ │⏱ 5h10m  │ │ │            │ │            │                        │\n'
'│  │ │Score: 45 │ │ │            │ │            │ WON (8) │ LOST (3)    │\n'
'│  │ └──────────┘ │ └────────────┘ └────────────┘                        │\n'
'│  └──────────────┘                                                        │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-06')
add_field_table(doc, [
    ('lead_search',          'Search Input',   'Tùy chọn', 'Max 100 ký tự',
     'Tìm theo tên, SĐT, email, model xe. Debounce 400ms. Highlight kết quả khớp.'),
    ('score_filter',         'Dropdown',       'Tùy chọn', 'Enum: ALL|HIGH(70+)|MEDIUM(40-69)|LOW(<40)',
     'Filter lead theo điểm. Score tính từ: mức độ tương tác, thời gian liên hệ, số lần trao đổi.'),
    ('kanban_column',        'Kanban Column',  'N/A', 'N/A',
     'Mỗi cột = 1 lead_status: NEW→CONTACTED→TEST_DRIVE_BOOKED→NEGOTIATING→WON/LOST. Manager thấy tất cả. Sale chỉ thấy lead của mình.'),
    ('lead_card',            'Draggable Card', 'N/A', 'N/A',
     'Drag sang cột kế tiếp = update lead_status. Không được kéo backward. AJAX update, rollback nếu lỗi (Optimistic UI).'),
    ('sla_countdown',        'Live Timer',     'N/A', 'N/A',
     'Đếm thời gian từ lần tương tác cuối (updated_at). Màu xanh < 8h. Vàng 8-24h. Đỏ > 24h (vi phạm SLA). Refresh 60s.'),
    ('lead_score',           'Badge',          'N/A', 'N/A',
     'Điểm 0-100. Màu xanh ≥70, vàng 40-69, đỏ <40. Tooltip: giải thích cách tính điểm.'),
    ('btn_quick_call',       'Icon Button 📞', 'N/A', 'N/A',
     'Giả lập cuộc gọi 1-click. Click → modal ghi chú cuộc gọi + update last_contact_at. Hiển thị call_count.'),
    ('btn_quick_paylink',    'Icon Button 💳', 'N/A', 'N/A',
     'Tạo link đặt cọc và gửi SMS/Zalo cho khách. Gọi POST /orders/{id}/send-paylink. Chỉ enable nếu khách đã có tài khoản.'),
    ('btn_add_lead',         'Primary Button', 'N/A', 'N/A',
     'Mở modal thêm lead mới. Nhập: tên, SĐT, email, model quan tâm, ghi chú. Gọi POST /crm/leads.'),
])

add_h3(doc, 'API Specs — SCR-06')
add_api_spec(doc, 'PATCH', '/api/v1/crm/leads/{lead_id}/status',
    auth='Bearer token. Role: SALE (lead của mình) | MANAGER (tất cả lead showroom).',
    desc='Cập nhật trạng thái lead. Chỉ cho phép transition forward theo State Machine.',
    path_params='lead_id: UUID — ID của crm_lead',
    request_body=(
        '{\n'
        '  "lead_status": "CONTACTED",    // NEW→CONTACTED→TEST_DRIVE_BOOKED→NEGOTIATING→WON|LOST\n'
        '  "lost_reason": null,           // required nếu status=LOST\n'
        '  "call_note": "Khách quan tâm Camry 2.5Q, hỏi thêm về trả góp"  // optional\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "lead_id": "lead-uuid-001",\n'
        '    "old_status": "NEW",\n'
        '    "new_status": "CONTACTED",\n'
        '    "updated_at": "2026-08-23T02:35:00+07:00"\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 422 Unprocessable (transition không hợp lệ)\n'
        '{ "success": false, "error": { "code": "ERR_LEAD_INVALID_TRANSITION", "message": "Không thể chuyển từ NEGOTIATING sang NEW" } }\n\n'
        'HTTP 403 Forbidden (BOLA)\n'
        '{ "success": false, "error": { "code": "ERR_LEAD_FORBIDDEN", "message": "Bạn không có quyền cập nhật lead này" } }\n\n'
        'HTTP 409 Conflict (lead đã bị re-assign)\n'
        '{ "success": false, "error": { "code": "ERR_LEAD_REASSIGNED", "message": "Lead này đã được giao cho người khác" } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/orders/{order_id}/send-paylink',
    auth='Bearer token. Role: SALE | MANAGER.',
    desc='Gửi link thanh toán qua SMS/Zalo tới khách hàng.',
    path_params='order_id: UUID — ID của order đang PENDING_PAYMENT',
    request_body=(
        '{\n'
        '  "channel": "SMS",     // SMS | ZALO_ZNS\n'
        '  "message_note": "Chào anh/chị, đây là link đặt cọc xe Camry..."  // optional custom msg\n'
        '}'
    ),
    happy_response='HTTP 200 OK\n{\n  "success": true,\n  "data": {\n    "message": "Link thanh toán đã được gửi thành công",\n    "channel": "SMS",\n    "sent_to": "090*****67"\n  }\n}',
    error_responses=(
        'HTTP 400 Bad Request (order không ở PENDING_PAYMENT)\n'
        '{ "success": false, "error": { "code": "ERR_ORDER_NOT_PENDING", "message": "Đơn hàng này không ở trạng thái chờ thanh toán" } }\n\n'
        'HTTP 429 Too Many Requests\n'
        '{ "success": false, "error": { "code": "ERR_PAYLINK_RATE_LIMIT", "message": "Đã gửi PayLink quá nhiều lần. Thử lại sau 10 phút." } }'
    )
)

doc.add_page_break()

# SCR-08 Inventory
add_h2(doc, '2.7. Luồng 8: Quản Lý Kho Xe & Manual VIN Hold (SCR-08 & SCR-12)')
add_bullet(doc, 'URL: /portal/inventory  |  Auth: MANAGER, ADMIN. Sale chỉ xem.')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [DEALER PORTAL]  Dashboard  [KHO XE & QUOTA]  Duyệt Giá  [👤 Manager] │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│  KHO XE & QUOTA   Showroom: Q1  [+ Nhập kho]  [🔄 Điều chuyển]         │\n'
'│                                                                          │\n'
'│ ┌────────┬───────────────┬──────┬────────┬────────────┬──────────────┐  │\n'
'│ │ VIN    │ Model         │ Màu  │ Status │ Quota      │ Hành động    │  │\n'
'│ ├────────┼───────────────┼──────┼────────┼────────────┼──────────────┤  │\n'
'│ │ VN1234 │ Camry 2.5Q    │Trắng │AVAILABLE│ 3/5 avail │ [Hold VIN]   │  │\n'
'│ │ VN5678 │ CR-V 1.5T     │ Đen  │VIN_HOLD │ Giữ: 23h │ [Xem] [Hủy] │  │\n'
'│ │ VN9012 │ Santa Fe 2.5  │Bạc   │SOLD    │ —         │ [Xem]        │  │\n'
'│ └────────┴───────────────┴──────┴────────┴────────────┴──────────────┘  │\n'
'│                                                                          │\n'
'│  [Nhập VIN] Modal:                                                       │\n'
'│  VIN *: [ vin_number               ]  ← ISO 3779 validation             │\n'
'│  Hold cho KH *: [ hold_customer_phone ] ← SĐT khách muốn giữ xe        │\n'
'│  Lý do: [ hold_reason              ]                                    │\n'
'│  [ XÁC NHẬN GIỮ XE 24H ]                                               │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-08')
add_field_table(doc, [
    ('showroom_filter',      'Dropdown',       'Tùy chọn', 'showroom_id. Admin thấy tất cả.',
     'Filter danh sách xe theo showroom. Manager chỉ thấy showroom của mình.'),
    ('vehicle_list',         'Data Table',     'N/A', 'N/A',
     'Danh sách xe với: VIN, model, màu, status, quota. Sortable theo cột. Filter theo status.'),
    ('vehicle_status',       'Status Badge',   'N/A', 'N/A',
     'AVAILABLE (xanh), VIN_HOLD (vàng, có countdown 24h), RESERVED (cam, soft-locked), SOLD (xám), TRANSIT (xanh nhạt).'),
    ('quota_display',        'Progress Bar',   'N/A', 'N/A',
     'Hiển thị available/total. Màu đỏ nếu available=0. Tooltip: "3 xe có thể đặt, 1 đang được giữ, 1 đã bán".'),
    ('btn_hold_vin',         'Action Button',  'N/A', 'N/A',
     'Chỉ Manager/Admin thấy. Mở modal nhập VIN, SĐT khách, lý do. Giữ xe 24h cho khách cụ thể.'),
    ('vin_number_input',     'Text Field',     'Bắt buộc', '^[A-HJ-NPR-Z0-9]{17}$. ISO 3779.',
     'Validate: không chứa I, O, Q. 17 ký tự chính xác. Backend verify checksum.'),
    ('hold_customer_phone',  'Text Field',     'Bắt buộc', '^(0|84)[35789][0-9]{8}$',
     'SĐT của khách muốn giữ xe. Lưu vào vin_hold_reservations.hold_customer_phone.'),
    ('hold_reason',          'Textarea',       'Tùy chọn', 'Max 500 ký tự',
     'Lý do giữ xe (khách đang cần thẩm định, sắp ký HĐ...). Ghi vào audit_log.'),
    ('btn_confirm_hold',     'Primary Button', 'N/A', 'N/A',
     'Gọi POST /inventory/manual-vin-hold. Atomic: increment soft_locked_count VÀ tạo vin_hold_reservations. Nếu quota=0 → lỗi.'),
])

add_h3(doc, 'API Specs — SCR-08')
add_api_spec(doc, 'POST', '/api/v1/inventory/manual-vin-hold',
    auth='Bearer token. Role: MANAGER | ADMIN. BOLA: chỉ hold xe tại showroom của mình (MANAGER).',
    desc='Giữ xe 24h cho khách cụ thể. Atomic: UPDATE vehicle_quotas + INSERT vin_hold_reservations trong một transaction.',
    request_body=(
        '{\n'
        '  "vin_number": "JN1FAAN15A0000001",\n'
        '  "hold_customer_phone": "0901234567",\n'
        '  "hold_reason": "Khach dang lam thu tuc ky HĐ, dự kiến chiều mai",\n'
        '  "showroom_id": "showroom-uuid-001"  // required cho Manager\n'
        '}'
    ),
    happy_response=(
        'HTTP 201 Created\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "reservation_id": "reservation-uuid-001",\n'
        '    "vin_number": "JN1FAAN15A0000001",\n'
        '    "hold_customer_phone": "090*****67",\n'
        '    "expires_at": "2026-08-24T02:35:00+07:00",\n'
        '    "remaining_quota": 2\n'
        '  }\n'
        '}\n'
        '// Side effects:\n'
        '// - soft_locked_count += 1 (atomic)\n'
        '// - Gửi SMS xác nhận tới hold_customer_phone\n'
        '// - Scheduler release sau 24h nếu không có đơn'
    ),
    error_responses=(
        'HTTP 409 Conflict (xe đã held hoặc quota=0)\n'
        '{ "success": false, "error": { "code": "ERR_VIN_QUOTA_FULL", "message": "Xe này đã hết quota. Không thể giữ thêm." } }\n\n'
        'HTTP 400 Bad Request (VIN không hợp lệ)\n'
        '{ "success": false, "error": { "code": "ERR_INVALID_VIN", "message": "Số VIN không đúng chuẩn ISO 3779" } }\n\n'
        'HTTP 403 Forbidden (SoD violation)\n'
        '{ "success": false, "error": { "code": "ERR_SOD_VIOLATION", "message": "Manager không có quyền Hold VIN tại showroom khác" } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/inventory/transfers/{transfer_id}/damage-report',
    auth='Bearer token. Role: MANAGER | ADMIN.',
    desc='Báo cáo xe bị hư hỏng trong quá trình chuyển kho. Rollback xe về showroom gốc.',
    path_params='transfer_id: UUID — ID của vehicle_transfer đang IN_TRANSIT',
    request_body=(
        '{\n'
        '  "damage_description": "Va cham nhe canh cua ben trai, xe tron ban phai",\n'
        '  "damage_photos": [\n'
        '    "https://cdn.autodealer.vn/damage/photo1.jpg"\n'
        '  ],\n'
        '  "logistics_loss_amount": 5000000  // ước tính thiệt hại logistics\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "transfer_id": "transfer-uuid-001",\n'
        '    "status": "TRANSIT_DAMAGED",\n'
        '    "vehicle_returned_to": "showroom-uuid-001",\n'
        '    "logistics_fee_adjusted": true\n'
        '  }\n'
        '}'
    ),
    error_responses='HTTP 400 Bad Request\n{ "success": false, "error": { "code": "ERR_TRANSFER_NOT_TRANSIT", "message": "Lệnh chuyển kho không ở trạng thái IN_TRANSIT" } }'
)

doc.add_page_break()

# SCR-09-CONF
add_h2(doc, '2.8. Luồng 10: Quản Trị Hệ Thống & Cấu Hình (SCR-09-CONF & SCR-13)')
add_bullet(doc, 'URL: /admin/configs | /admin/audit-logs  |  Auth: ADMIN only.')

add_wireframe(doc,
'┌─────────────────────────────────────────────────────────────────────────┐\n'
'│  [ADMIN PORTAL]  [CẤU HÌNH HỆ THỐNG]  Users  Audit Logs  [👤 Admin]    │\n'
'├─────────────────────────────────────────────────────────────────────────┤\n'
'│  CẤU HÌNH THAM SỐ HỆ THỐNG                          [Lưu tất cả 💾]   │\n'
'│                                                                          │\n'
'│  ┌─────────────────────────────────────────────────────────────────┐    │\n'
'│  │ Thanh toán & Đặt cọc                                             │    │\n'
'│  │ Tiền cọc mặc định (VNĐ):    [  50,000,000            ]          │    │\n'
'│  │ Timeout soft-lock (phút):   [  15                    ]          │    │\n'
'│  │ Tỷ lệ vay tối đa (%):       [  80                    ]          │    │\n'
'│  ├─────────────────────────────────────────────────────────────────┤    │\n'
'│  │ Chiết khấu & Hoa hồng                                            │    │\n'
'│  │ Giới hạn chiết khấu Manager (%): [ 10                ]          │    │\n'
'│  │ Tích lũy chiết khấu tối đa/đơn:  [ 30,000,000        ]          │    │\n'
'│  ├─────────────────────────────────────────────────────────────────┤    │\n'
'│  │ Hoàn tiền & SLA                                                   │    │\n'
'│  │ SLA hoàn tiền (ngày):       [  3                     ]          │    │\n'
'│  └─────────────────────────────────────────────────────────────────┘    │\n'
'│                                                                          │\n'
'│  NHẬT KÝ KIỂM TOÁN (SCR-13)                                            │\n'
'│  Lọc: [Action ▾] [Entity ▾] [User ▾] [Từ ngày 📅] [Đến ngày 📅]      │\n'
'│  [🔍 Tìm Correlation ID]                                                │\n'
'│                                                                          │\n'
'│ ┌────────────┬────────┬───────────┬───────────┬────────────────────┐  │\n'
'│ │ Thời gian  │ Actor  │ Action    │ Entity    │ Chi tiết / Decrypt │  │\n'
'│ ├────────────┼────────┼───────────┼───────────┼────────────────────┤  │\n'
'│ │ 02:31:00   │ Admin  │ PASSWORD_ │ users/    │ [Xem] [🔓 Decrypt] │  │\n'
'│ │            │        │ RESET     │ uuid-001  │                    │  │\n'
'│ └────────────┴────────┴───────────┴───────────┴────────────────────┘  │\n'
'│                                                                          │\n'
'└─────────────────────────────────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-09-CONF')
add_field_table(doc, [
    ('cfg_default_deposit',     'Number Input',  'Bắt buộc', 'Min 10,000,000. Max 500,000,000. Số nguyên.',
     'Tiền cọc mặc định toàn hệ thống (VNĐ). Lưu vào system_configs với key=cfg_default_deposit_amount.'),
    ('cfg_soft_lock_timeout',   'Number Input',  'Bắt buộc', 'Min 5. Max 60. Số nguyên (phút).',
     'Thời gian checkout hết hạn. Scheduler dùng giá trị này để cancel đơn PENDING_PAYMENT quá hạn.'),
    ('cfg_max_loan_percent',    'Number Input',  'Bắt buộc', 'Min 50. Max 90. Số nguyên %.',
     'Tỷ lệ vay tối đa trên giá xe. Client-side validation trên SCR-04 (loan calculator) cũng dùng giá trị này.'),
    ('cfg_manager_discount',    'Number Input',  'Bắt buộc', 'Min 0. Max 20. Số thập phân %.',
     'Giới hạn chiết khấu Manager có thể duyệt. Vượt ngưỡng → tự động escalate Admin.'),
    ('cfg_max_discount_total',  'Number Input',  'Bắt buộc', 'Min 0. Max 100,000,000. Số nguyên VNĐ.',
     'Tổng chiết khấu tích lũy tối đa/đơn hàng. Vượt ngưỡng → PENDING_ADMIN.'),
    ('cfg_refund_sla_days',     'Number Input',  'Bắt buộc', 'Min 1. Max 30. Số nguyên (ngày).',
     'SLA hoàn tiền. Cron job escalation khi payout_due_date < TODAY.'),
    ('btn_save_all',            'Primary Button','N/A', 'N/A',
     'Gọi PUT /system/configs với tất cả configs đã thay đổi trong session. Ghi audit_log với old_value/new_value. Confirm dialog trước khi save.'),
    ('filter_audit_action',     'Dropdown',      'Tùy chọn', 'String enum từ distinct action values',
     'Lọc log theo loại hành động (LOGIN, LOGOUT, PASSWORD_RESET, ORDER_CREATE, REFUND_APPROVE...).'),
    ('filter_audit_entity',     'Dropdown',      'Tùy chọn', 'String enum entity_type',
     'Lọc log theo loại đối tượng: users, orders, payments, vehicles, discount_requests...'),
    ('filter_correlation_id',   'Text Search',   'Tùy chọn', 'UUID hoặc partial string',
     'Tìm toàn bộ log cùng correlation_id (trace 1 request end-to-end). Hữu ích debug.'),
    ('btn_decrypt_pii',         'Danger Button', 'N/A', 'N/A',
     'Chỉ Admin thấy. Click → confirm dialog + lý do decrypt. Gọi POST /audit/decrypt-pii. Ghi audit log: DECRYPT_PII. Hiện CCCD plaintext 30 giây rồi ẩn lại.'),
])

add_h3(doc, 'API Specs — SCR-09-CONF & SCR-13')
add_api_spec(doc, 'GET', '/api/v1/system/configs',
    auth='Bearer token. Role: ADMIN only.',
    desc='Lấy toàn bộ cấu hình hệ thống từ bảng system_configs.',
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": [\n'
        '    {\n'
        '      "config_key": "cfg_default_deposit_amount",\n'
        '      "config_value": "50000000",\n'
        '      "data_type": "INTEGER",\n'
        '      "description": "Số tiền đặt cọc mặc định (VNĐ)",\n'
        '      "updated_by": { "id": "admin-uuid", "name": "Super Admin" },\n'
        '      "updated_at": "2026-08-20T10:00:00+07:00"\n'
        '    },\n'
        '    {\n'
        '      "config_key": "cfg_soft_lock_timeout_minutes",\n'
        '      "config_value": "15",\n'
        '      "data_type": "INTEGER",\n'
        '      "description": "Thời gian soft-lock hết hạn (phút)"\n'
        '    }\n'
        '  ]\n'
        '}'
    ),
    error_responses='HTTP 403 Forbidden\n{ "success": false, "error": { "code": "ERR_FORBIDDEN", "message": "Chỉ Super Admin mới có quyền truy cập cấu hình hệ thống" } }'
)

add_api_spec(doc, 'PUT', '/api/v1/system/configs',
    auth='Bearer token. Role: ADMIN only.',
    desc='Cập nhật một hoặc nhiều tham số cấu hình. Ghi audit log với old_value.',
    request_body=(
        '{\n'
        '  "configs": [\n'
        '    { "config_key": "cfg_default_deposit_amount", "config_value": "60000000" },\n'
        '    { "config_key": "cfg_soft_lock_timeout_minutes", "config_value": "20" }\n'
        '  ]\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "updated": 2,\n'
        '    "configs": [\n'
        '      { "config_key": "cfg_default_deposit_amount", "old_value": "50000000", "new_value": "60000000" },\n'
        '      { "config_key": "cfg_soft_lock_timeout_minutes", "old_value": "15", "new_value": "20" }\n'
        '    ]\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 400 Bad Request (giá trị ngoài range)\n'
        '{ "success": false, "error": { "code": "ERR_CONFIG_INVALID", "message": "Giá trị không hợp lệ", "details": [{"key":"cfg_default_deposit_amount","msg":"Giá trị phải từ 10,000,000 đến 500,000,000"}] } }\n\n'
        'HTTP 403 Forbidden\n'
        '{ "success": false, "error": { "code": "ERR_FORBIDDEN", "message": "Chỉ Super Admin mới có quyền sửa cấu hình" } }'
    )
)

add_api_spec(doc, 'GET', '/api/v1/audit/logs',
    auth='Bearer token. Role: ADMIN.',
    desc='Lấy danh sách audit logs với filter và pagination.',
    query_params=(
        '?action=PASSWORD_RESET      // filter by action\n'
        '&entity_type=users          // filter by entity type\n'
        '&actor_user_id=uuid         // filter by actor\n'
        '&from_date=2026-08-01       // date range start\n'
        '&to_date=2026-08-23         // date range end\n'
        '&correlation_id=abc123      // trace specific request\n'
        '&page=1&limit=50'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": [\n'
        '    {\n'
        '      "id": "log-uuid-001",\n'
        '      "actor_type": "USER",\n'
        '      "actor": { "id": "user-uuid", "name": "Admin", "role": "ADMIN" },\n'
        '      "action": "PASSWORD_RESET",\n'
        '      "entity_type": "users",\n'
        '      "entity_id": "user-uuid-002",\n'
        '      "old_value": { "password_hash": "[MASKED]" },\n'
        '      "new_value": { "password_hash": "[MASKED]" },\n'
        '      "ip_address": "192.168.1.100",\n'
        '      "correlation_id": "req-abc123",\n'
        '      "created_at": "2026-08-23T02:31:00+07:00",\n'
        '      "has_pii": false\n'
        '    }\n'
        '  ],\n'
        '  "meta": { "page": 1, "limit": 50, "total": 1240 }\n'
        '}'
    ),
    error_responses='HTTP 403 Forbidden\n{ "success": false, "error": { "code": "ERR_FORBIDDEN", "message": "Chỉ Admin mới có quyền xem Audit Logs" } }'
)

add_api_spec(doc, 'POST', '/api/v1/audit/decrypt-pii',
    auth='Bearer token. Role: ADMIN only. MFA required.',
    desc='Giải mã dữ liệu PII (CCCD). Bắt buộc ghi audit_log với lý do. Chỉ hiện 30 giây.',
    request_body=(
        '{\n'
        '  "entity_type": "customer_profiles",\n'
        '  "entity_id": "profile-uuid-001",\n'
        '  "decrypt_reason": "Xác minh danh tính khách hàng để giải quyết khiếu nại #TKT-001"\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "identity_card_number": "001200001234",  // Plaintext, TTL 30s\n'
        '    "auto_mask_after_seconds": 30\n'
        '  }\n'
        '}\n'
        '// Side effects: ghi audit_log action=DECRYPT_PII với decrypted_user_ids=[profile-uuid-001]'
    ),
    error_responses=(
        'HTTP 403 Forbidden (không phải Admin)\n'
        '{ "success": false, "error": { "code": "ERR_FORBIDDEN", "message": "Không có quyền decrypt PII" } }\n\n'
        'HTTP 400 Bad Request (thiếu lý do)\n'
        '{ "success": false, "error": { "code": "ERR_DECRYPT_REASON_REQUIRED", "message": "Phải nhập lý do giải mã" } }'
    )
)

doc.add_page_break()

# Refund (SCR-07/SCR-11)
add_h2(doc, '2.9. Luồng 7: Hoàn Tiền Cọc Đa Cấp (SCR-11 → SCR-07)')
add_bullet(doc, 'URL: /portal/my-orders/{id}  |  Auth: CUSTOMER xem SCR-11, Manager/Admin xem SCR-07')

add_wireframe(doc,
'┌──────────────────────────────────┬──────────────────────────────────────┐\n'
'│  YÊU CẦU HOÀN TIỀN CỌC (SCR-11) │  DUYỆT HOÀN CỌC OVERRIDE (SCR-07)   │\n'
'├──────────────────────────────────┼──────────────────────────────────────┤\n'
'│ Mã ĐH: ORD-2026-001             │ Đơn: ORD-2026-001  Cọc: 50,000,000   │\n'
'│ Lý do: [ Vay bị từ chối    ▾ ] │                                      │\n'
'│                                  │ KH yêu cầu: BANK_LOAN_REJECTED      │\n'
'│ Chứng từ từ chối vay:           │ Thư từ chối: [Xem PDF]               │\n'
'│ [📎 Upload file PDF] (required)  │                                      │\n'
'│                                  │ Manager Override:                    │\n'
'│ Thông tin STK nhận tiền:        │ ☐ Bỏ qua yêu cầu chứng từ          │\n'
'│ Số TK: [ bank_account_number ] │ Lý do: [ override_reason      ]      │\n'
'│ Ngân hàng: [ bank_name      ] │                                      │\n'
'│ Tên TK: [ bank_account_name ] │ [✅ DUYỆT HOÀN CỌC]                 │\n'
'│                                  │ [❌ TỪ CHỐI]                        │\n'
'│ [ GỬI YÊU CẦU HOÀN TIỀN  ]    │ [📤 Giao Admin xử lý]               │\n'
'└──────────────────────────────────┴──────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-11 (Refund Request)')
add_field_table(doc, [
    ('refund_reason_type',     'Dropdown',      'Bắt buộc', 'Enum: BANK_LOAN_REJECTED|SYSTEM_TIMEOUT_ERROR|FORCE_MAJEURE',
     'Lý do hoàn cọc. BANK_LOAN_REJECTED → yêu cầu bank_rejection_letter_url. FORCE_MAJEURE → Manager Override.'),
    ('bank_rejection_letter',  'File Upload',   'Bắt buộc nếu BANK_LOAN_REJECTED', 'PDF. Max 10MB.',
     'Upload file thư từ chối vay từ ngân hàng. Lưu S3, URL vào refund_requests.bank_rejection_letter_url.'),
    ('bank_account_number',    'Text Field',    'Bắt buộc', '9-20 chữ số',
     'STK nhận tiền hoàn cọc. Pre-fill từ customer_profiles nếu có. Validate format số tài khoản VN.'),
    ('bank_name',              'Dropdown/Text', 'Bắt buộc', 'Tên ngân hàng hợp lệ',
     'Dropdown danh sách ngân hàng VN hoặc nhập tay. Dùng để chuyển khoản manual.'),
    ('bank_account_name',      'Text Field',    'Bắt buộc', 'Min 2, Max 100 ký tự. Uppercase.',
     'Tên chủ tài khoản (phải khớp với tên trên sổ ngân hàng). Uppercase tự động.'),
    ('btn_submit_refund',      'Primary Button','N/A', 'N/A',
     'Gọi POST /refunds/request. Disabled nếu thiếu required fields. Sau submit → order hiện trạng thái "Đang chờ xử lý hoàn tiền".'),
    ('manager_override_check', 'Checkbox (Mgr)','N/A', 'N/A',
     'Chỉ Manager thấy. Cho phép bỏ qua yêu cầu chứng từ. Bắt buộc nhập override_reason nếu tick.'),
    ('override_reason',        'Textarea (Mgr)','Bắt buộc nếu override', 'Min 20, Max 500 ký tự',
     'Lý do Manager Override. Ghi vào refund_requests.manager_override_reason và audit_log.'),
    ('btn_approve_refund',     'Success Button','N/A', 'N/A',
     'Manager/Admin duyệt. Gọi POST /refunds/{id}/decision {action: APPROVE}. Tạo payout_due_date = TODAY + cfg_refund_sla_days.'),
    ('btn_reject_refund',      'Danger Button', 'N/A', 'N/A',
     'Từ chối hoàn cọc. Gọi POST /refunds/{id}/decision {action: REJECT, reason: ...}. Order → DEPOSIT_PAID.'),
])

add_h3(doc, 'API Specs — SCR-07/11 Refund')
add_api_spec(doc, 'POST', '/api/v1/refunds/request',
    auth='Bearer token. Role: CUSTOMER (self) | SALE (hộ customer).\nBOLA: order.customer_id phải == JWT.sub HOẶC role >= SALE.',
    desc='Tạo yêu cầu hoàn cọc. Order phải ở trạng thái DEPOSIT_PAID hoặc sau đó.',
    request_body=(
        '{\n'
        '  "order_id": "order-uuid-001",\n'
        '  "refund_reason_type": "BANK_LOAN_REJECTED",\n'
        '  "bank_rejection_letter_url": "https://cdn.autodealer.vn/docs/rejection.pdf",\n'
        '  "bank_account_number": "0123456789",\n'
        '  "bank_account_name": "NGUYEN VAN A",\n'
        '  "bank_name": "Vietcombank"\n'
        '  // Bổ sung v23: requested_by_user_id tự lấy từ JWT.sub\n'
        '  // requester_role tự set từ JWT.role\n'
        '}'
    ),
    happy_response=(
        'HTTP 201 Created\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "refund_id": "refund-uuid-001",\n'
        '    "refund_code": "REF-20260823-001",\n'
        '    "status": "PENDING_MANAGER",\n'
        '    "payout_due_date": "2026-08-26",\n'
        '    "message": "Yêu cầu hoàn tiền đã được gửi. Showroom sẽ xử lý trong 3 ngày làm việc."\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 400 Bad Request (thiếu chứng từ vay)\n'
        '{ "success": false, "error": { "code": "ERR_MISSING_REJECTION_DOCUMENT", "message": "Cần đính kèm thư từ chối vay từ ngân hàng" } }\n\n'
        'HTTP 409 Conflict (đã có refund request cho đơn này)\n'
        '{ "success": false, "error": { "code": "ERR_DUPLICATE_REFUND", "message": "Đơn hoàn cọc cho đơn hàng này đã được tạo trước đó" } }\n\n'
        'HTTP 400 Bad Request (order không ở trạng thái đủ điều kiện)\n'
        '{ "success": false, "error": { "code": "ERR_ORDER_INVALID_STATUS", "message": "Đơn hàng phải ở trạng thái DEPOSIT_PAID hoặc sau đó để yêu cầu hoàn cọc" } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/refunds/{refund_id}/decision',
    auth='Bearer token. Role: MANAGER (PENDING_MANAGER) | ADMIN (PENDING_ADMIN).',
    desc='Duyệt hoặc từ chối yêu cầu hoàn cọc.',
    path_params='refund_id: UUID — ID của refund_request',
    request_body=(
        '{\n'
        '  "action": "APPROVE",           // APPROVE | REJECT\n'
        '  "reject_reason": null,         // required nếu action=REJECT\n'
        '  "manager_override_reason": "Khach hang thuoc trường hợp bất khả kháng, cho phép hoàn cọc không cần chứng từ"\n'
        '                               // required nếu Manager Override\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "refund_id": "refund-uuid-001",\n'
        '    "status": "PENDING_ADMIN",  // hoặc COMPLETED nếu Admin approve\n'
        '    "next_step": "Chờ Admin xét duyệt giải ngân",\n'
        '    "payout_due_date": "2026-08-26"\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 403 Forbidden\n'
        '{ "success": false, "error": { "code": "ERR_REFUND_FORBIDDEN", "message": "Bạn không có quyền duyệt yêu cầu hoàn tiền này" } }\n\n'
        'HTTP 400 Bad Request\n'
        '{ "success": false, "error": { "code": "ERR_REFUND_ALREADY_PROCESSED", "message": "Đơn hoàn tiền này đã được xử lý" } }'
    )
)

doc.add_page_break()

# Trade-in
add_h2(doc, '2.10. Luồng 6: Thu Cũ Đổi Mới (Trade-in) & Cấn Trừ Dòng Tiền')
add_bullet(doc, 'URL: Form trên SCR-02 / SCR-04 / SCR-11  |  Auth: CUSTOMER | SALE (hộ)')

add_wireframe(doc,
'┌────────────────────────────┬─────────────────────────────────────────────┐\n'
'│ FORM THU CŨ ĐỔI MỚI (SCR-02) │ TRẠNG THÁI THẨM ĐỊNH (SCR-11)           │\n'
'├────────────────────────────┼─────────────────────────────────────────────┤\n'
'│ Hãng xe cũ *: [Toyota ▾]  │ Xe cũ: Toyota Vios E 2020 — 62,000 km       │\n'
'│ Dòng xe *: [Vios ▾]        │ Trạng thái: 🟡 INSPECTION_PENDING           │\n'
'│ Phiên bản: [E 1.5MT ▾]     │                                             │\n'
'│ Năm SX *: [2020 ▾]         │ Sau thẩm định:                             │\n'
'│ Màu: [Trắng ▾]             │ Giá thẩm định: 320,000,000 đ               │\n'
'│                             │ Giá kỳ vọng bạn: 350,000,000 đ            │\n'
'│ Số km đã đi *:             │ Chênh lệch: -30,000,000 đ                  │\n'
'│ [ odo_km         ] km      │                                             │\n'
'│ Warning: ⚠️ > 300,000 km    │ [✅ Tôi đồng ý - Cấn trừ ngay]            │\n'
'│                             │ [❌ Từ chối thẩm định]                     │\n'
'│ Giá kỳ vọng *:             │                                             │\n'
'│ [ expected_price  ] đ      │ Nếu đồng ý: final_price sẽ giảm 320M      │\n'
'│                             │                                             │\n'
'│ Ghi chú tình trạng:        │ INSPECTION_FAILED:                         │\n'
'│ [ condition_notes ]        │ ❌ Xe bị thủy kích / tai nạn nặng           │\n'
'│                             │ Cấn trừ = 0, Cash cần thêm = 320M         │\n'
'│ [GỬI YÊU CẦU THẨM ĐỊNH]  │ [Xem lại options]                          │\n'
'└────────────────────────────┴─────────────────────────────────────────────┘')

add_h3(doc, 'Field Spec Table — SCR-02/11 Trade-in')
add_field_table(doc, [
    ('brand_old_car',     'Dropdown',       'Bắt buộc', 'brand_id UUID',
     'Hãng xe cũ. Fetch từ /catalog/brands. Khi chọn → fetch models phù hợp.'),
    ('model_old_car',     'Dropdown',       'Bắt buộc', 'model_id UUID',
     'Dòng xe cũ (ví dụ: Vios). Phụ thuộc vào brand đã chọn.'),
    ('variant_old_car',   'Dropdown',       'Tùy chọn', 'string',
     'Phiên bản cụ thể (ví dụ: 1.5MT E). Nhập text nếu không có trong list.'),
    ('year_old_car',      'Dropdown',       'Bắt buộc', 'Integer, 1990 đến năm hiện tại',
     'Năm sản xuất. Tự động tính tuổi xe = CurrentYear - year. Hiện tuổi xe bên cạnh.'),
    ('color_old_car',     'Dropdown',       'Tùy chọn', 'string',
     'Màu xe cũ. Ảnh hưởng giá thẩm định (màu đặc biệt có thể cộng điểm).'),
    ('odo_km',            'Number Input',   'Bắt buộc', 'Min 0. Max 999,999. Số nguyên km.',
     'Số km đã đi. Warning icon + text nếu > 300,000 km. Ảnh hưởng giá thẩm định.'),
    ('expected_price',    'Number Input',   'Bắt buộc', 'Min 10,000,000đ. Max = listed_price xe mới.',
     'Giá kỳ vọng của khách. Hiển thị format số có dấu phẩy. Nếu khách kỳ vọng quá cao → hiện gợi ý giá hợp lý.'),
    ('condition_notes',   'Textarea',       'Tùy chọn', 'Max 1000 ký tự',
     'Mô tả tình trạng xe: tai nạn, sửa chữa lớn, điểm xước... Giúp Appraiser thẩm định chính xác hơn.'),
    ('condition_photos',  'File Upload',    'Tùy chọn', 'JPG/PNG. Max 5 ảnh. 5MB/ảnh.',
     'Upload ảnh xe cũ (4 góc + nội thất). Lưu file_attachments. Preview thumbnails.'),
    ('btn_submit_ti',     'Primary Button', 'N/A', 'N/A',
     'Gọi POST /trade-in/submit. Tạo trade_in_request với status PENDING_INSPECTION. Gửi thông báo Sale/Manager.'),
    ('btn_accept_price',  'Success Button', 'N/A', 'N/A',
     'Chỉ hiện sau khi Appraiser cập nhật appraised_price. Gọi PATCH /trade-in/{id}/status {status: ACCEPTED}. Trigger auto-update final_price trên order.'),
    ('btn_reject_price',  'Danger Button',  'N/A', 'N/A',
     'Từ chối giá thẩm định. Gọi PATCH /trade-in/{id}/status {status: REJECTED}. Xóa trade_in_credit_value khỏi order.'),
])

add_h3(doc, 'API Specs — Trade-in')
add_api_spec(doc, 'POST', '/api/v1/trade-in/submit',
    auth='Bearer token. Role: CUSTOMER | SALE.',
    desc='Gửi yêu cầu thẩm định xe cũ. Liên kết với order_id nếu đã có đơn.',
    request_body=(
        '{\n'
        '  "order_id": "order-uuid-001",   // optional: link tới đơn đặt cọc\n'
        '  "brand_name": "Toyota",\n'
        '  "model_name": "Vios",\n'
        '  "variant_name": "1.5MT E",      // optional\n'
        '  "year": 2020,\n'
        '  "color": "Trắng",\n'
        '  "odo_km": 62000,\n'
        '  "expected_price": 350000000,\n'
        '  "condition_notes": "Xe còn mới, chưa tai nạn, 1 chủ",\n'
        '  "photo_urls": ["https://cdn.../ti1.jpg"]\n'
        '}'
    ),
    happy_response=(
        'HTTP 201 Created\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "trade_in_id": "ti-uuid-001",\n'
        '    "status": "PENDING_INSPECTION",\n'
        '    "message": "Yêu cầu thẩm định đã được ghi nhận. Nhân viên sẽ liên hệ trong 24h."\n'
        '  }\n'
        '}'
    ),
    error_responses='HTTP 400 Bad Request\n{ "success": false, "error": { "code": "ERR_TI_INVALID", "message": "Giá kỳ vọng vượt quá giá niêm yết xe mới" } }'
)

add_api_spec(doc, 'PATCH', '/api/v1/trade-in/{trade_in_id}/status',
    auth='Bearer token. Role: SALE | MANAGER | CUSTOMER (chỉ ACCEPTED/REJECTED).',
    desc='Cập nhật trạng thái trade-in. Khi ACCEPTED → auto-update orders.final_price.',
    path_params='trade_in_id: UUID',
    request_body=(
        '{\n'
        '  "status": "ACCEPTED",          // PENDING_INSPECTION|INSPECTION_DONE|ACCEPTED|REJECTED|INSPECTION_FAILED\n'
        '  "appraised_price": 320000000,  // required nếu status=INSPECTION_DONE\n'
        '  "appraiser_notes": "Xe tốt, một vài vết xước nhỏ"\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "trade_in_id": "ti-uuid-001",\n'
        '    "status": "ACCEPTED",\n'
        '    "appraised_price": 320000000,\n'
        '    "order_updated": {\n'
        '      "order_id": "order-uuid-001",\n'
        '      "old_final_price": 1290000000,\n'
        '      "new_final_price": 970000000,\n'
        '      "trade_in_credit_value": 320000000\n'
        '    }\n'
        '  }\n'
        '}\n'
        '// Side effects khi ACCEPTED:\n'
        '// - UPDATE orders SET trade_in_credit_value = 320M\n'
        '// - Recalculate final_price = listed_price - trade_in_credit_value\n'
        '// - Ghi audit_log\n'
        '// Side effects khi INSPECTION_FAILED:\n'
        '// - trade_in_credit_value = 0\n'
        '// - Recalculate final_price (bỏ phần cấn trừ)'
    ),
    error_responses=(
        'HTTP 400 Bad Request (transition không hợp lệ)\n'
        '{ "success": false, "error": { "code": "ERR_TI_INVALID_TRANSITION", "message": "Không thể chuyển sang trạng thái này" } }'
    )
)

doc.add_page_break()

# Bank Switch (SCR-11)
add_h2(doc, '2.11. Luồng 5: Mua Xe Trả Góp & Bank Switch (SCR-11)')
add_h3(doc, 'API Specs — Loan & Bank Switch')
add_api_spec(doc, 'POST', '/api/v1/loans/{loan_id}/switch-bank',
    auth='Bearer token. Role: CUSTOMER.\nBOLA: loan.order.customer_id phải == JWT.sub.',
    desc='Đổi ngân hàng xét duyệt vay. Chỉ được thực hiện sau REJECTED hoặc PARTIALLY_APPROVED. Max 3 lần.',
    path_params='loan_id: UUID — ID của loan_application',
    request_body=(
        '{\n'
        '  "new_bank_name": "BIDV",\n'
        '  "reuse_existing_docs": true,     // true: copy URLs từ hồ sơ cũ, không re-upload\n'
        '  "new_financial_docs_urls": [],   // required nếu reuse_existing_docs=false\n'
        '  "co_borrower_change": false      // có đổi người vay phụ không\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "new_loan_id": "loan-uuid-002",\n'
        '    "bank_name": "BIDV",\n'
        '    "status": "SUBMITTED",\n'
        '    "switch_count": 2,\n'
        '    "remaining_switches": 1,\n'
        '    "order_status": "BANK_APPROVING"\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 400 Bad Request (đã hết lượt đổi)\n'
        '{ "success": false, "error": { "code": "ERR_BANK_SWITCH_LIMIT", "message": "Bạn đã dùng hết 3 lượt đổi ngân hàng" } }\n\n'
        'HTTP 400 Bad Request (loan chưa bị reject)\n'
        '{ "success": false, "error": { "code": "ERR_BANK_SWITCH_NOT_ALLOWED", "message": "Chỉ có thể đổi ngân hàng sau khi bị từ chối hoặc duyệt thiếu" } }'
    )
)

add_api_spec(doc, 'POST', '/api/v1/mock/bank-decision',
    auth='Internal sandbox endpoint. Rate limit: 10 req/min.',
    desc='Mock endpoint giả lập quyết định của ngân hàng. Chỉ dùng trong sandbox/development.',
    request_body=(
        '{\n'
        '  "loan_id": "loan-uuid-001",\n'
        '  "decision": "APPROVED",          // APPROVED | PARTIALLY_APPROVED | REJECTED\n'
        '  "approved_amount": 800000000,    // required nếu APPROVED hoặc PARTIALLY_APPROVED\n'
        '  "rejection_reason": null,        // required nếu REJECTED\n'
        '  "simulated_delay_ms": 1500       // mock delay\n'
        '}'
    ),
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "loan_id": "loan-uuid-001",\n'
        '    "decision": "APPROVED",\n'
        '    "order_status_updated_to": "BANK_APPROVED",\n'
        '    "sse_event_sent": true\n'
        '  }\n'
        '}'
    ),
    error_responses='HTTP 403 Forbidden (dùng trên production)\n{ "success": false, "error": { "code": "ERR_SANDBOX_ONLY", "message": "Endpoint này chỉ khả dụng trên môi trường sandbox" } }'
)

add_api_spec(doc, 'GET', '/api/v1/orders/{order_id}',
    auth='Bearer token. BOLA: CUSTOMER chỉ xem order của mình. SALE/MANAGER/ADMIN xem theo phân quyền.',
    desc='Lấy chi tiết đơn hàng bao gồm tất cả thông tin liên quan.',
    path_params='order_id: UUID',
    happy_response=(
        'HTTP 200 OK\n'
        '{\n'
        '  "success": true,\n'
        '  "data": {\n'
        '    "id": "order-uuid-001",\n'
        '    "order_code": "ORD-20260823-0001",\n'
        '    "status": "BANK_APPROVING",\n'
        '    "purchase_type": "AUTO_LOAN",\n'
        '    "vehicle": { "name": "Toyota Camry 2.5Q", "vin": "JN1...", "color": "Trắng" },\n'
        '    "showroom": { "name": "AutoDealership Q1", "phone": "028-..." },\n'
        '    "pricing": {\n'
        '      "listed_price": 1290000000,\n'
        '      "accessories_total": 5000000,\n'
        '      "insurance_total": 12000000,\n'
        '      "trade_in_credit": 320000000,\n'
        '      "credit_applied": 10000000,\n'
        '      "deposit_paid": 50000000,\n'
        '      "final_price": 977000000\n'
        '    },\n'
        '    "loan": {\n'
        '      "bank": "Techcombank",\n'
        '      "requested_amount": 800000000,\n'
        '      "status": "IN_REVIEW",\n'
        '      "switch_count": 0\n'
        '    },\n'
        '    "status_history": [\n'
        '      { "status": "PENDING_PAYMENT", "at": "2026-08-23T02:30:00+07:00" },\n'
        '      { "status": "DEPOSIT_PAID", "at": "2026-08-23T02:31:00+07:00" },\n'
        '      { "status": "BANK_APPROVING", "at": "2026-08-23T02:32:00+07:00" }\n'
        '    ],\n'
        '    "created_at": "2026-08-23T02:30:00+07:00"\n'
        '  }\n'
        '}'
    ),
    error_responses=(
        'HTTP 404 Not Found\n{ "success": false, "error": { "code": "ERR_ORDER_NOT_FOUND", "message": "Không tìm thấy đơn hàng" } }\n\n'
        'HTTP 403 Forbidden (BOLA)\n{ "success": false, "error": { "code": "ERR_ORDER_FORBIDDEN", "message": "Bạn không có quyền xem đơn hàng này" } }'
    )
)

doc.add_page_break()

# ───────────────────────────────────
# SECTION 3: ORDER STATE MACHINE
# ───────────────────────────────────
add_h1(doc, '3. Order State Machine — Tất Cả Transitions Hợp Lệ (B2.1 Fix)')

add_body(doc,
    'Bất kỳ transition nào KHÔNG có trong bảng dưới đây đều bị CẤM. '
    'Backend phải validate transition trước khi UPDATE orders.status. '
    'Mọi transition phải ghi vào order_status_history.',
    bold=False)

add_state_machine_table(doc, [
    ('PENDING_PAYMENT',         'DEPOSIT_PAID',           'Payment success webhook',          'received_amount >= deposit_amount',                 'PAYMENT_GATEWAY'),
    ('PENDING_PAYMENT',         'PAYMENT_FAILED',         'Payment failed webhook',            'Gateway response: failed/timeout',                  'PAYMENT_GATEWAY'),
    ('PENDING_PAYMENT',         'CANCELED',               'Timeout 15 phút',                   'Scheduler: created_at < NOW()-15min',               'SCHEDULER'),
    ('PENDING_PAYMENT',         'CANCELED',               'Customer cancel',                   'User action — có confirm dialog',                   'USER'),
    ('PAYMENT_FAILED',          'PENDING_PAYMENT',        'User retry payment',                'attempt_no <= 3',                                   'USER'),
    ('PAYMENT_FAILED',          'CANCELED',               'Max retries exceeded',              'attempt_no > 3 → auto-cancel',                      'SCHEDULER'),
    ('DEPOSIT_PAID',            'BANK_APPROVING',         'Submit loan application',           'purchase_type = AUTO_LOAN',                         'USER'),
    ('DEPOSIT_PAID',            'PROCESSING',             'Direct purchase confirm',           'purchase_type = DIRECT, Sale/Manager confirm',      'USER / SYSTEM'),
    ('DEPOSIT_PAID',            'REFUND_REQUESTED',       'Customer request refund',           'Tạo bản ghi refund_request',                        'USER'),
    ('BANK_APPROVING',          'BANK_APPROVED',          'Bank full approval',                'approved_amount = requested_loan_amount',            'PAYMENT_GATEWAY'),
    ('BANK_APPROVING',          'BANK_PARTIALLY_APPROVED','Bank partial approval',             'approved_amount < requested_loan_amount',           'PAYMENT_GATEWAY'),
    ('BANK_APPROVING',          'BANK_REJECTED',          'Bank rejection',                    'Bank trả về rejection',                             'PAYMENT_GATEWAY'),
    ('BANK_APPROVED',           'PROCESSING',             'Auto-transition',                   'Immediately sau khi bank approve full',             'SYSTEM'),
    ('BANK_PARTIALLY_APPROVED', 'PROCESSING',             'Customer accepts partial',          'Customer xác nhận bù tiền mặt phần thiếu',         'USER'),
    ('BANK_PARTIALLY_APPROVED', 'BANK_APPROVING',         'Switch bank',                       'switch_count < 3',                                  'USER'),
    ('BANK_REJECTED',           'BANK_APPROVING',         'Switch bank',                       'switch_count < 3',                                  'USER'),
    ('BANK_REJECTED',           'REFUND_REQUESTED',       'All banks failed',                  'switch_count >= 3 → tự động suggest refund',        'SYSTEM / USER'),
    ('PROCESSING',              'READY_FOR_DELIVERY',     'VIN assigned + PDI complete',       'vehicles.vin_number NOT NULL AND pdi_completed=true', 'SALE / MANAGER'),
    ('READY_FOR_DELIVERY',      'DELIVERED',              'Delivery confirmed',                'Sale/Manager confirm + ký biên bản bàn giao',       'SALE / MANAGER'),
    ('DELIVERED',               'COMPLETED',              'All post-delivery done',            'registration + PDI cert + full payment confirmed', 'SYSTEM / MANAGER'),
    ('REFUND_REQUESTED',        'REFUNDED',               'Refund approved + disbursed',       'Admin approval + bank transfer confirmed',           'ADMIN'),
    ('REFUND_REQUESTED',        'DEPOSIT_PAID',           'Refund rejected',                   'Admin/Manager từ chối → order active lại',          'ADMIN / MANAGER'),
    ('REFUNDED',                'CANCELED',               'Auto-transition final',             'Sau khi REFUNDED → tự chuyển CANCELED',             'SYSTEM'),
])

add_body(doc,
    '⚠️ DELIVERED vs COMPLETED: DELIVERED = bàn giao vật lý (ký biên bản). '
    'COMPLETED = hoàn tất mọi thủ tục (đăng ký xe, PDI certificate, thanh toán đủ 100%). '
    'COMPLETED là trạng thái cuối cùng của đơn hàng thành công.',
    color=C_WARN, bold=True)

doc.add_page_break()

# ───────────────────────────────────
# SECTION 4: ATOMIC QUOTA SQL
# ───────────────────────────────────
add_h1(doc, '4. Atomic Quota Engine — SQL Spec Chi Tiết (B3.1 Fix)')

add_body(doc,
    'Giải quyết race condition khi 100 users checkout đồng thời chiếc xe cuối cùng. '
    'Sử dụng PostgreSQL row-level locking + optimistic locking với version column.')

add_h2(doc, '4.1. Atomic Soft-Lock SQL')
add_code_block(doc,
'-- Bước 1: Atomic increment với double guard\n'
'-- Cả hai điều kiện phải true: còn quota VÀ version đúng (optimistic locking)\n'
'UPDATE vehicle_quotas\n'
'SET\n'
'    soft_locked_count = soft_locked_count + 1,\n'
'    version = version + 1\n'
'WHERE\n'
'    id = $1                                      -- quota_id\n'
'    AND soft_locked_count < total_physical_count  -- còn chỗ\n'
'    AND version = $2                              -- tránh race condition\n'
'RETURNING id, soft_locked_count, total_physical_count, version;\n'
'\n'
'-- Nếu 0 rows returned → Race condition hoặc hết xe → Return ERR_QUOTA_EXCEEDED\n'
'-- Nếu 1 row returned → Lock thành công → Tiến hành tạo order')

add_h2(doc, '4.2. Safety Net — Recalculate khi hiển thị quota')
add_code_block(doc,
'-- E4.3 Fix: soft_locked_count có thể bị stale nếu scheduler chưa chạy\n'
'-- Luôn recount từ orders table khi HIỂN THỊ quota cho user\n'
'SELECT\n'
'    vq.total_physical_count,\n'
'    vq.soft_locked_count AS stored_locked,\n'
'    COUNT(o.id) AS actual_locked,\n'
'    vq.total_physical_count - COUNT(o.id) AS display_available\n'
'FROM vehicle_quotas vq\n'
'LEFT JOIN orders o ON\n'
'    o.variant_id = vq.variant_id\n'
'    AND o.showroom_id = vq.showroom_id\n'
'    AND o.status = \'PENDING_PAYMENT\'\n'
'    AND o.created_at > NOW() - INTERVAL \'15 minutes\'  -- trong thời gian soft-lock\n'
'WHERE vq.id = $1\n'
'GROUP BY vq.id, vq.total_physical_count, vq.soft_locked_count;')

add_h2(doc, '4.3. Scheduler Release Expired Locks (chạy mỗi 60s)')
add_code_block(doc,
'-- B3.2 Fix: Giải phóng soft_locked_count khi đơn hàng PENDING_PAYMENT hết hạn 15 phút\n'
'BEGIN;\n'
'\n'
'-- Hủy đơn hàng quá hạn\n'
'UPDATE orders\n'
'SET status = \'CANCELED\'\n'
'WHERE status = \'PENDING_PAYMENT\'\n'
'  AND created_at < NOW() - INTERVAL \'15 minutes\';\n'
'\n'
'-- Giải phóng quota cho từng variant/showroom\n'
'UPDATE vehicle_quotas vq\n'
'SET soft_locked_count = GREATEST(0,\n'
'    vq.soft_locked_count - expired.expired_count\n'
')\n'
'FROM (\n'
'    SELECT variant_id, showroom_id, COUNT(*) as expired_count\n'
'    FROM orders\n'
'    WHERE status = \'CANCELED\'\n'
'      AND updated_at >= NOW() - INTERVAL \'65 seconds\'  -- vừa cancel trong vòng scheduler\n'
'    GROUP BY variant_id, showroom_id\n'
') expired\n'
'WHERE vq.variant_id = expired.variant_id\n'
'  AND vq.showroom_id = expired.showroom_id;\n'
'\n'
'COMMIT;')

doc.add_page_break()

# ───────────────────────────────────
# SECTION 5: DB — DDL FIXES & NEW TABLES
# ───────────────────────────────────
add_h1(doc, '5. Database Schema Updates v24 — DDL Fixes & Bảng Mới')

add_h2(doc, '5.1. Sửa lỗi DDL (E1.x)')
add_code_block(doc,
'-- E1.1 CRITICAL FIX: notifications.user_id thiếu kiểu dữ liệu\n'
'-- DDL gốc sẽ FAIL khi chạy. Sửa:\n'
'ALTER TABLE notifications\n'
'    ALTER COLUMN user_id SET DATA TYPE UUID USING user_id::UUID,\n'
'    ALTER COLUMN user_id SET NOT NULL,\n'
'    ADD CONSTRAINT fk_notifications_user FOREIGN KEY (user_id) REFERENCES users(id);\n'
'\n'
'-- E1.2 HIGH FIX: vehicle_models UNIQUE constraint\n'
'ALTER TABLE vehicle_models\n'
'    ADD CONSTRAINT unique_brand_model UNIQUE(brand_id, name);\n'
'\n'
'-- E1.3 HIGH FIX: crm_leads phone uniqueness\n'
'ALTER TABLE crm_leads\n'
'    ADD CONSTRAINT unique_lead_phone_variant UNIQUE(phone, interested_variant_id);\n'
'\n'
'-- E1.4/1.5/1.6 MEDIUM FIX: Foreign keys to orders\n'
'ALTER TABLE discount_requests\n'
'    ADD CONSTRAINT fk_discount_order FOREIGN KEY (order_id) REFERENCES orders(id);\n'
'ALTER TABLE trade_in_requests\n'
'    ADD CONSTRAINT fk_trade_in_order FOREIGN KEY (order_id) REFERENCES orders(id);\n'
'ALTER TABLE vin_hold_reservations\n'
'    ADD CONSTRAINT fk_vin_hold_order FOREIGN KEY (order_id) REFERENCES orders(id);\n'
'\n'
'-- E4.10 LOW FIX: email_verified column\n'
'ALTER TABLE users ADD COLUMN IF NOT EXISTS email_verified BOOLEAN DEFAULT FALSE;\n'
'\n'
'-- E4.8 MEDIUM FIX: soft-delete pattern\n'
'ALTER TABLE orders ADD COLUMN IF NOT EXISTS deleted_at TIMESTAMPTZ NULL;\n'
'ALTER TABLE payments ADD COLUMN IF NOT EXISTS deleted_at TIMESTAMPTZ NULL;\n'
'-- NOTE: Không hard-delete orders/payments/audit_logs bao giờ\n'
'\n'
'-- E4.7 MEDIUM FIX: outbox retry cap\n'
'ALTER TABLE outbox_events ADD COLUMN IF NOT EXISTS max_retries INTEGER DEFAULT 5;\n'
'ALTER TABLE outbox_events ADD COLUMN IF NOT EXISTS dead_lettered_at TIMESTAMPTZ NULL;\n'
'-- Khi retry_count >= max_retries: status → DEAD_LETTER + alert admin\n'
'\n'
'-- E4.6 MEDIUM FIX: GIN index for specs_json\n'
'CREATE INDEX IF NOT EXISTS idx_variant_specs ON vehicle_variants USING gin(specs_json);')

add_h2(doc, '5.2. Bảng mới (E2.x)')
add_code_block(doc,
'-- E2.1 CRITICAL: system_configs\n'
'CREATE TABLE IF NOT EXISTS system_configs (\n'
'    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
'    config_key VARCHAR(100) UNIQUE NOT NULL,\n'
'    config_value TEXT NOT NULL,\n'
'    data_type VARCHAR(20) NOT NULL CHECK (data_type IN (\'STRING\',\'INTEGER\',\'DECIMAL\',\'BOOLEAN\',\'JSON\')),\n'
'    description TEXT,\n'
'    updated_by UUID REFERENCES users(id),\n'
'    updated_at TIMESTAMPTZ DEFAULT NOW(),\n'
'    created_at TIMESTAMPTZ DEFAULT NOW()\n'
');\n'
'-- Seed:\n'
'INSERT INTO system_configs (config_key, config_value, data_type, description) VALUES\n'
'    (\'cfg_default_deposit_amount\', \'50000000\', \'INTEGER\', \'Số tiền đặt cọc mặc định (VNĐ)\'),\n'
'    (\'cfg_soft_lock_timeout_minutes\', \'15\', \'INTEGER\', \'Thời gian soft-lock hết hạn (phút)\'),\n'
'    (\'cfg_manager_discount_limit_percent\', \'10\', \'DECIMAL\', \'Giới hạn chiết khấu Manager (%)\'),\n'
'    (\'cfg_refund_sla_days\', \'3\', \'INTEGER\', \'SLA hoàn tiền (ngày)\'),\n'
'    (\'cfg_max_loan_amount_percent\', \'80\', \'DECIMAL\', \'Tỷ lệ vay tối đa (% giá xe)\'),\n'
'    (\'cfg_max_discount_accumulation\', \'30000000\', \'INTEGER\', \'Tổng chiết khấu tích lũy max/đơn (VNĐ)\')\n'
'ON CONFLICT (config_key) DO NOTHING;\n'
'\n'
'-- E2.2 CRITICAL: test_drive_bookings\n'
'CREATE TABLE IF NOT EXISTS test_drive_bookings (\n'
'    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
'    slot_id UUID NOT NULL REFERENCES test_drive_slots(id),\n'
'    customer_user_id UUID REFERENCES users(id) NULL,\n'
'    customer_name VARCHAR(100) NOT NULL,\n'
'    customer_phone VARCHAR(15) NOT NULL,\n'
'    driver_license VARCHAR(12) NULL,\n'
'    is_on_behalf BOOLEAN DEFAULT FALSE,\n'
'    on_behalf_customer_name VARCHAR(100) NULL,\n'
'    on_behalf_customer_phone VARCHAR(15) NULL,\n'
'    status VARCHAR(30) NOT NULL DEFAULT \'CONFIRMED\'\n'
'        CHECK (status IN (\'CONFIRMED\',\'ARRIVED\',\'COMPLETED\',\'NO_SHOW\',\'CANCELLED\')),\n'
'    notes TEXT NULL,\n'
'    created_at TIMESTAMPTZ DEFAULT NOW(),\n'
'    updated_at TIMESTAMPTZ DEFAULT NOW()\n'
');\n'
'\n'
'-- E2.5 HIGH: user_sessions\n'
'CREATE TABLE IF NOT EXISTS user_sessions (\n'
'    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
'    user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,\n'
'    refresh_token_hash VARCHAR(255) NOT NULL UNIQUE,\n'
'    device_info TEXT,\n'
'    ip INET,\n'
'    user_agent TEXT,\n'
'    expires_at TIMESTAMPTZ NOT NULL,\n'
'    created_at TIMESTAMPTZ DEFAULT NOW()\n'
');\n'
'CREATE INDEX idx_user_sessions_user ON user_sessions(user_id);\n'
'\n'
'-- E2.6 CRITICAL: password_reset_tokens\n'
'CREATE TABLE IF NOT EXISTS password_reset_tokens (\n'
'    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
'    user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,\n'
'    token_hash VARCHAR(255) NOT NULL UNIQUE,\n'
'    expires_at TIMESTAMPTZ NOT NULL,\n'
'    is_used BOOLEAN DEFAULT FALSE,\n'
'    ip_requested INET,\n'
'    created_at TIMESTAMPTZ DEFAULT NOW()\n'
');\n'
'\n'
'-- E2.3 HIGH: accessories_catalog\n'
'CREATE TABLE IF NOT EXISTS accessories_catalog (\n'
'    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
'    name VARCHAR(200) NOT NULL,\n'
'    sku VARCHAR(50) UNIQUE,\n'
'    price NUMERIC(15,2) NOT NULL CHECK (price >= 0),\n'
'    category VARCHAR(100),\n'
'    description TEXT,\n'
'    image_url TEXT,\n'
'    is_active BOOLEAN DEFAULT TRUE,\n'
'    created_at TIMESTAMPTZ DEFAULT NOW(),\n'
'    updated_at TIMESTAMPTZ DEFAULT NOW()\n'
');\n'
'\n'
'-- E2.4 MEDIUM: insurance_products\n'
'CREATE TABLE IF NOT EXISTS insurance_products (\n'
'    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
'    name VARCHAR(200) NOT NULL,\n'
'    provider VARCHAR(100) NOT NULL,\n'
'    coverage_type VARCHAR(50) NOT NULL\n'
'        CHECK (coverage_type IN (\'COMPREHENSIVE\',\'THIRD_PARTY\',\'BODY\',\'FIRE_THEFT\')),\n'
'    annual_price NUMERIC(15,2) NOT NULL CHECK (annual_price > 0),\n'
'    is_active BOOLEAN DEFAULT TRUE,\n'
'    created_at TIMESTAMPTZ DEFAULT NOW()\n'
');\n'
'\n'
'-- E2.7 HIGH: file_attachments\n'
'CREATE TABLE IF NOT EXISTS file_attachments (\n'
'    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
'    uploaded_by UUID NOT NULL REFERENCES users(id),\n'
'    entity_type VARCHAR(50) NOT NULL\n'
'        CHECK (entity_type IN (\'ORDER\',\'REFUND\',\'LOAN\',\'TRADE_IN\',\'PROFILE\')),\n'
'    entity_id UUID NOT NULL,\n'
'    file_url TEXT NOT NULL,\n'
'    file_size INTEGER,\n'
'    mime_type VARCHAR(100),\n'
'    original_filename VARCHAR(255),\n'
'    created_at TIMESTAMPTZ DEFAULT NOW()\n'
');\n'
'CREATE INDEX idx_attachments_entity ON file_attachments(entity_type, entity_id);\n'
'\n'
'-- E2.8 LOW: showroom_operating_hours\n'
'CREATE TABLE IF NOT EXISTS showroom_operating_hours (\n'
'    id UUID PRIMARY KEY DEFAULT gen_random_uuid(),\n'
'    showroom_id UUID NOT NULL REFERENCES showrooms(id) ON DELETE CASCADE,\n'
'    day_of_week INT NOT NULL CHECK (day_of_week BETWEEN 0 AND 6), -- 0=Sunday\n'
'    open_time TIME NOT NULL,\n'
'    close_time TIME NOT NULL,\n'
'    is_closed BOOLEAN DEFAULT FALSE,\n'
'    UNIQUE(showroom_id, day_of_week)\n'
');')

add_h2(doc, '5.3. Missing Indexes (E3.x)')
add_code_block(doc,
'-- E3.1 HIGH: Sale query đơn hàng\n'
'CREATE INDEX IF NOT EXISTS idx_orders_sale_id ON orders(sale_id);\n'
'\n'
'-- E3.2 HIGH: Manager filter by showroom\n'
'CREATE INDEX IF NOT EXISTS idx_orders_showroom_id ON orders(showroom_id);\n'
'\n'
'-- E3.3 HIGH: JOIN payments khi load order\n'
'CREATE INDEX IF NOT EXISTS idx_payments_order_id ON payments(order_id);\n'
'\n'
'-- E3.4 HIGH: JOIN loan info khi load order\n'
'CREATE INDEX IF NOT EXISTS idx_loan_applications_order_id ON loan_applications(order_id);\n'
'\n'
'-- E3.5 MEDIUM: Lookup refund cho order\n'
'CREATE INDEX IF NOT EXISTS idx_refund_requests_order_id ON refund_requests(order_id);\n'
'\n'
'-- E3.6/E3.7 MEDIUM: Vehicle transfers\n'
'CREATE INDEX IF NOT EXISTS idx_transfers_from_showroom ON vehicle_transfers(from_showroom_id);\n'
'CREATE INDEX IF NOT EXISTS idx_transfers_to_showroom ON vehicle_transfers(to_showroom_id);\n'
'\n'
'-- E3.8 HIGH: Time-range queries on audit_logs\n'
'CREATE INDEX IF NOT EXISTS idx_audit_logs_created_at ON audit_logs(created_at);\n'
'\n'
'-- E3.9 MEDIUM: Audit trail per entity\n'
'CREATE INDEX IF NOT EXISTS idx_audit_logs_entity ON audit_logs(entity_type, entity_id);\n'
'\n'
'-- E3.10 MEDIUM: Kanban filter\n'
'CREATE INDEX IF NOT EXISTS idx_crm_leads_status ON crm_leads(lead_status);\n'
'\n'
'-- E3.11 HIGH: Inventory query by showroom\n'
'CREATE INDEX IF NOT EXISTS idx_vehicles_showroom_id ON vehicles(showroom_id);\n'
'\n'
'-- E3.12 MEDIUM: Order status timeline\n'
'CREATE INDEX IF NOT EXISTS idx_order_status_history_order ON order_status_history(order_id);\n'
'\n'
'-- E3.13 MEDIUM: Credit balance history\n'
'CREATE INDEX IF NOT EXISTS idx_credit_tx_account_time ON credit_transactions(account_id, created_at);\n'
'\n'
'-- E3.14 MEDIUM: Notification list per user\n'
'CREATE INDEX IF NOT EXISTS idx_notifications_user_time ON notifications(user_id, created_at);')

add_h2(doc, '5.4. Triggers & Data Integrity (E4.x)')
add_code_block(doc,
'-- E4.4: auto-update updated_at trigger\n'
'CREATE OR REPLACE FUNCTION set_updated_at() RETURNS TRIGGER AS $$\n'
'BEGIN NEW.updated_at = NOW(); RETURN NEW; END;\n'
'$$ LANGUAGE plpgsql;\n'
'\n'
'-- Apply to all tables with updated_at:\n'
'DO $$ DECLARE t TEXT;\n'
'BEGIN\n'
'    FOR t IN SELECT unnest(ARRAY[\n'
'        \'users\',\'orders\',\'payments\',\'vehicles\',\'customer_profiles\',\n'
'        \'crm_leads\',\'accessories_catalog\',\'test_drive_bookings\',\'vehicle_quotas\'\n'
'    ]) LOOP\n'
'        EXECUTE format(\n'
'            \'CREATE TRIGGER trg_%s_updated_at BEFORE UPDATE ON %s\n'
'             FOR EACH ROW EXECUTE FUNCTION set_updated_at();\', t, t);\n'
'    END LOOP;\n'
'END $$;\n'
'\n'
'-- E4.1: vehicle_quotas sync trigger\n'
'CREATE OR REPLACE FUNCTION sync_vehicle_quota_count() RETURNS TRIGGER AS $$\n'
'BEGIN\n'
'    IF TG_OP = \'INSERT\' THEN\n'
'        UPDATE vehicle_quotas SET total_physical_count = total_physical_count + 1\n'
'        WHERE variant_id = NEW.variant_id AND showroom_id = NEW.showroom_id;\n'
'    ELSIF TG_OP = \'DELETE\' THEN\n'
'        UPDATE vehicle_quotas SET total_physical_count = GREATEST(0, total_physical_count - 1)\n'
'        WHERE variant_id = OLD.variant_id AND showroom_id = OLD.showroom_id;\n'
'    ELSIF TG_OP = \'UPDATE\' AND OLD.showroom_id != NEW.showroom_id THEN\n'
'        UPDATE vehicle_quotas SET total_physical_count = GREATEST(0, total_physical_count - 1)\n'
'        WHERE variant_id = OLD.variant_id AND showroom_id = OLD.showroom_id;\n'
'        UPDATE vehicle_quotas SET total_physical_count = total_physical_count + 1\n'
'        WHERE variant_id = NEW.variant_id AND showroom_id = NEW.showroom_id;\n'
'    END IF;\n'
'    RETURN COALESCE(NEW, OLD);\n'
'END;\n'
'$$ LANGUAGE plpgsql;\n'
'\n'
'CREATE TRIGGER trg_vehicles_quota_sync\n'
'    AFTER INSERT OR DELETE OR UPDATE OF showroom_id ON vehicles\n'
'    FOR EACH ROW EXECUTE FUNCTION sync_vehicle_quota_count();')

doc.add_page_break()

# ───────────────────────────────────
# SECTION 6: VALIDATION RULES
# ───────────────────────────────────
add_h1(doc, '6. Validation Rules Updates (C1-C2)')

add_generic_table(doc,
    ['Finding', 'Field / Screen', 'Mức độ', 'Rule Cũ (Sai)', 'Rule Mới (Đúng)'],
    [
        ('C1.1', 'identity_input SĐT', '🟡 HIGH',
         '^(0|84)[3|5|7|8|9][0-9]{8}$\n❌ Bug: | trong [] là ký tự pipe, không phải OR',
         '^(0|84)[35789][0-9]{8}$\n✅ Character class [35789]'),
        ('C1.2', 'password (tất cả forms)', '🟠 MEDIUM',
         'Min 8 ký tự, 1 hoa, 1 số',
         'Min 10 ký tự, 1 hoa, 1 thường, 1 số, 1 ký tự đặc biệt (@$!%*?&)\n^(?=.*[A-Z])(?=.*[a-z])(?=.*\\d)(?=.*[@$!%*?&])[A-Za-z\\d@$!%*?&]{10,}$'),
        ('C1.3', 'identity_card_number (CCCD)', '🟠 MEDIUM',
         '12 chữ số (chỉ validate độ dài)',
         '12 chữ số + structural check:\n- 3 số đầu = mã tỉnh (001-096)\n- Số 4 = giới tính + thế kỷ sinh (0,1,2,3)'),
        ('C1.4', 'identity_card_date', '🔵 LOW',
         'Date <= Today (cho phép 1900)',
         'Date >= 2012-01-01 AND Date <= Today\n(CCCD mới được cấp từ 2012)'),
        ('C1.5', 'otp_digits', '🟠 MEDIUM',
         'Sandbox hardcode 888888, không rõ production',
         'Sandbox: 888888 cố định.\nProduction: 6 số random, gửi SMS brandname, TTL 60s.\nTách logic bằng ENV=sandbox'),
        ('C1.6', 'vin_number (SCR-08)', '🟡 HIGH',
         'VARCHAR(17) không validate format',
         '^[A-HJ-NPR-Z0-9]{17}$\n(ISO 3779: loại I,O,Q)\nBackend validate checksum digit (vị trí 9)'),
        ('C1.7', 'email (update profile)', '🟠 MEDIUM',
         'Cập nhật email trực tiếp, không verify',
         'Sau update: gửi verification link.\nemail_verified=false cho đến khi click link.\nKhông cho login bằng email mới cho đến khi verified'),
        ('C1.8', 'odo_km (Trade-in)', '🔵 LOW',
         'Không có range validation',
         'Min: 0. Max: 999,999 km.\nWarning UI nếu > 300,000 km: "Xe chạy nhiều có thể ảnh hưởng giá thẩm định"'),
        ('C1.9', 'expected_price (Trade-in)', '🔵 LOW',
         'Không có range validation',
         'Min: 10,000,000 đ.\nMax: listed_price của xe mới muốn mua.\nHiện gợi ý giá thị trường nếu nhập quá cao'),
        ('C2.1', 'Deposit amount (SCR-04)', '🟡 HIGH',
         'UI cố định 50M nhưng DB có min_deposit_amount per variant — mâu thuẫn',
         'Option B: Deposit linh hoạt >= min_deposit_amount.\nUI: "Đặt cọc tối thiểu: {min}đ"\nDefault = 50M từ system_configs'),
        ('C2.3', 'Loan cap (SCR-11)', '🟠 MEDIUM',
         'Không có client-side validation',
         'Client: requested_loan_amount <= 0.8 * (final_price - deposit)\nDisable [Nộp đơn] nếu vượt. Hiển thị: "Tối đa vay 80% giá xe (sau cọc)"'),
        ('C2.4', 'Discount tích lũy (SCR-07)', '🟡 HIGH',
         'TC_QA_005 test nhưng logic chưa spec rõ: per order hay per customer?',
         'Per order: SUM(discount_amount) WHERE order_id=$1 AND status IN (APPROVED, PENDING)\nNếu > 30,000,000đ → tự động PENDING_ADMIN'),
    ]
)

doc.add_page_break()

# ───────────────────────────────────
# SECTION 7: NOTIFICATION MATRIX
# ───────────────────────────────────
add_h1(doc, '7. Notification Matrix — Đầy Đủ (B5.1 Fix)')

add_body(doc, 'Spec đầy đủ: Event nào → Channel nào → Template nào → Thời điểm nào.')

add_notif_table(doc, [
    ('PAYMENT_SUCCESS',           'IN_APP + SMS',           'Customer',      'TPL_PAY_SUCCESS',     'Ngay khi webhook nhận DEPOSIT_PAID'),
    ('PAYMENT_FAILED',            'IN_APP + SMS',           'Customer',      'TPL_PAY_FAILED',      'Ngay khi webhook nhận PAYMENT_FAILED'),
    ('PAYMENT_EXPIRED',           'IN_APP',                 'Customer',      'TPL_PAY_EXPIRED',     'Khi scheduler cancel đơn (15p timeout)'),
    ('ORDER_CANCELED',            'IN_APP + SMS',           'Customer',      'TPL_ORDER_CANCELED',  'Khi order → CANCELED'),
    ('BANK_APPROVED',             'IN_APP + EMAIL',         'Customer',      'TPL_BANK_APPROVED',   'Ngay khi bank approve full'),
    ('BANK_PARTIALLY_APPROVED',   'IN_APP + EMAIL + SMS',   'Customer',      'TPL_BANK_PARTIAL',    'Ngay khi bank approve partial'),
    ('BANK_REJECTED',             'IN_APP + EMAIL + SMS',   'Customer',      'TPL_BANK_REJECTED',   'Ngay khi bank reject'),
    ('REFUND_REQUESTED',          'IN_APP',                 'Sale + Manager','TPL_REFUND_NEW',      'Khi customer tạo refund request'),
    ('REFUND_APPROVED',           'IN_APP + SMS',           'Customer',      'TPL_REFUND_APPROVED', 'Khi Admin approve refund'),
    ('REFUND_REJECTED',           'IN_APP + SMS',           'Customer',      'TPL_REFUND_REJECTED', 'Khi Admin/Manager reject refund'),
    ('REFUND_SLA_OVERDUE',        'IN_APP + EMAIL',         'Admin + Manager','TPL_REFUND_OVERDUE', 'Cron daily: payout_due_date < TODAY'),
    ('TEST_DRIVE_CONFIRMED',      'SMS + IN_APP',           'Customer',      'TPL_TD_CONFIRM',      'Ngay khi booking confirmed'),
    ('TEST_DRIVE_REMINDER',       'SMS',                    'Customer',      'TPL_TD_REMINDER',     '2 giờ trước slot_start (scheduled job)'),
    ('TEST_DRIVE_NO_SHOW',        'IN_APP',                 'Sale',          'TPL_TD_NO_SHOW',      'Khi Sale đánh dấu NO_SHOW'),
    ('ORDER_READY_DELIVERY',      'IN_APP + SMS',           'Customer',      'TPL_READY_DELIVERY',  'Khi status → READY_FOR_DELIVERY'),
    ('ORDER_DELIVERED',           'IN_APP + EMAIL',         'Customer',      'TPL_DELIVERED',       'Khi Sale/Manager confirm delivery'),
    ('ORDER_COMPLETED',           'IN_APP + EMAIL',         'Customer',      'TPL_COMPLETED',       'Khi tất cả thủ tục hoàn tất'),
    ('DISCOUNT_PENDING_ADMIN',    'IN_APP + EMAIL',         'Admin',         'TPL_DISC_REVIEW',     'Khi discount > cfg_max_discount'),
    ('DISCOUNT_APPROVED',         'IN_APP',                 'Customer + Sale','TPL_DISC_APPROVED',  'Khi Admin/Manager approve discount'),
    ('LEAD_ASSIGNED',             'IN_APP',                 'Sale',          'TPL_LEAD_ASSIGNED',   'Khi Manager assign lead cho Sale'),
    ('LEAD_SLA_WARNING',          'IN_APP',                 'Sale + Manager','TPL_LEAD_SLA',        'Mỗi 4h nếu lead NEW chưa CONTACTED'),
    ('VIN_HOLD_EXPIRING',         'IN_APP + SMS',           'Manager',       'TPL_VIN_EXPIRING',    '1 giờ trước khi Hold VIN hết hạn'),
    ('VIN_HOLD_RELEASED',         'IN_APP',                 'Manager',       'TPL_VIN_RELEASED',    'Khi Hold VIN hết hạn, auto-release'),
])

doc.add_page_break()

# ───────────────────────────────────
# SECTION 8: NFR & QA
# ───────────────────────────────────
add_h1(doc, '8. Tiêu Chuẩn Phi Chức Năng (NFR) & QA Automation')

add_h2(doc, '8.1. Performance Standards')
add_generic_table(doc,
    ['Hạng mục', 'Tiêu chuẩn', 'Phương pháp kiểm thử'],
    [
        ('Core Web Vitals', 'LCP < 1.5s, CLS = 0, FID < 100ms', 'Google Lighthouse CI + PageSpeed Insights trên Preview'),
        ('API Response Time', 'P99 < 200ms tại 500 RPM đồng thời', 'k6 Load Test: 500 VU × 5 phút'),
        ('DB Query Time', 'P95 < 50ms cho tất cả queries', 'pg_stat_statements monitoring'),
        ('SSE Latency', 'Event delivery < 500ms sau webhook', 'Manual + k6 WebSocket test'),
        ('File Upload', 'Timeout 30s, max 10MB', 'Playwright E2E test upload'),
    ]
)

add_h2(doc, '8.2. Security Standards')
add_generic_table(doc,
    ['Hạng mục', 'Yêu cầu'],
    [
        ('PII Encryption', 'AES-256-GCM cho CCCD. Masking 100% trên UI/API (001200001234 → 001200****34)'),
        ('JWT Security', 'RS256 algorithm. Access token TTL 15p. Refresh token rotation với reuse detection.'),
        ('Rate Limiting', 'Upstash Redis Token Bucket. Login: 5 req/min. OTP: 3 req/hour. Checkout: 10 req/day/user.'),
        ('BOLA Prevention', 'Mỗi resource endpoint verify: resource.owner_id == JWT.sub OR role >= MANAGER'),
        ('Webhook Security', 'HMAC-SHA256 signature verification. Idempotency key để tránh duplicate processing.'),
        ('CORS Policy', 'Whitelist: production domain + staging. Không dùng wildcard *'),
    ]
)

add_h2(doc, '8.3. QA Test Scenarios (Updated v24)')
add_generic_table(doc,
    ['Mã TC', 'Tên kịch bản', 'Steps & Test Data', 'Expected Result'],
    [
        ('TC_001', 'Đăng ký tài khoản mới', 'POST /auth/register với full_name, email, phone, password đủ mạnh, OTP 888888', 'HTTP 201. User tạo thành công. access_token trả về. Email verify gửi đi.'),
        ('TC_002', 'Đăng ký trùng SĐT/email', 'POST /auth/register với phone/email đã tồn tại', 'HTTP 409 ERR_REG_001. Không tạo user mới.'),
        ('TC_003', 'Login bị khóa 30p', 'Login sai password 5 lần liên tiếp', 'Lần 5: HTTP 423 ERR_AUTH_002. Account bị lock. Unlock sau 30p.'),
        ('TC_004', 'Checkout đồng thời quota=1', 'Bắn đồng thời 2 POST /orders/deposit khi quota=1', 'Chỉ 1 request thành công (201). Request còn lại: 409 ERR_QUOTA_EXCEEDED.'),
        ('TC_005', 'Order State Machine violation', 'Cố update order CANCELED → DEPOSIT_PAID', 'HTTP 422. Transition bị chặn. order_status_history không thay đổi.'),
        ('TC_006', 'Discount tích lũy > 30M', 'Tạo discount request 25M, sau đó tạo thêm 20M cùng đơn', 'Request 1: PENDING_MANAGER. Request 2: Tổng 45M > 30M → PENDING_ADMIN tự động.'),
        ('TC_007', 'Trade-in INSPECTION_FAILED', 'Appraiser cập nhật status=INSPECTION_FAILED', 'trade_in_credit_value=0. final_price recalculated. Không cấn trừ. Audit log ghi nhận.'),
        ('TC_008', 'Refund Manager Override', 'Manager tick override, nhập lý do, approve', 'Refund status PENDING_ADMIN. manager_override_reason ghi vào audit_log.'),
        ('TC_009', 'VIN Hold + Checkout cùng lúc quota=1', 'Manual VIN Hold + Checkout đồng thời', 'Atomic lock: chỉ 1 thành công. Người kia nhận ERR_VIN_QUOTA_FULL.'),
        ('TC_010', 'Password Reset token reuse', 'Dùng reset token đã dùng rồi', 'HTTP 400 ERR_FP_TOKEN_USED. Không cho phép đặt lại mật khẩu lần 2.'),
        ('TC_011', 'Webhook HMAC invalid signature', 'POST /payments/mock-webhook với X-Webhook-Signature sai', 'HTTP 401 ERR_WEBHOOK_SIGNATURE. Không update order status.'),
        ('TC_012', 'Bank switch > 3 lần', 'Cố switch bank lần thứ 4', 'HTTP 400 ERR_BANK_SWITCH_LIMIT. switch_count vẫn là 3.'),
    ]
)

doc.add_page_break()

# ───────────────────────────────────
# SECTION 9: SPRINT PLAN
# ───────────────────────────────────
add_h1(doc, '9. Lộ Trình Triển Khai 5 Sprint (v24 Updated)')

add_generic_table(doc,
    ['Sprint', 'Mục tiêu', 'Deliverables'],
    [
        ('Pre-Sprint 1\n(Review Gate)',
         'Giải quyết 18 CRITICAL — Không bắt đầu Sprint 1 khi chưa hoàn tất',
         '✅ SCR-00-REG Đăng ký (UI + API + DB)\n✅ SCR-00-FP Quên mật khẩu (UI + API + DB)\n✅ Order State Machine table documented\n✅ DDL bugs fixed (notifications, vehicle_models, FKs)\n✅ Missing tables: system_configs, test_drive_bookings, password_reset_tokens\n✅ API request body schemas cho tất cả endpoints\n✅ BOLA authorization documented per-endpoint'),
        ('Sprint 1',
         'Design System, Core Auth, Database Foundation',
         '- Next.js 14 App Router + Design Tokens\n- DDL v24 complete (tất cả new tables + indexes + triggers)\n- SCR-00 Login + OTP + Register + ForgotPassword fully working\n- user_sessions refresh token rotation\n- Outbox Pattern + Credit Ledger\n- Email verification flow\n- Error code chuẩn hóa (ERR_CLIENT_xxx / ERR_SERVER_xxx)'),
        ('Sprint 2',
         'Catalog, Quota Engine, Test Drive',
         '- Master Data: brands, models, variants, accessories, insurance\n- SCR-01 Danh mục + filters + Lifestyle Quiz\n- SCR-02 PDP + 360° Viewer + Loan Calculator\n- SCR-03 Đặt lịch lái thử (test_drive_bookings)\n- Atomic Quota SQL (SELECT FOR UPDATE + scheduler)\n- showroom_operating_hours integration\n- Mobile responsive wireframes (SCR-00, SCR-01, SCR-02, SCR-03)'),
        ('Sprint 3',
         'Checkout, Payment, SSE Realtime',
         '- SCR-04 Checkout (accessories, insurance, credit wallet)\n- Mock Payment Simulator (VietQR + VNPay)\n- SSE Realtime Payment Listener\n- SCR-10 Kết quả thanh toán (4 states)\n- Payment retry flow (max 3, PAYMENT_FAILED → PENDING_PAYMENT)\n- Webhook HMAC signature verification\n- Bank Switch wizard (SCR-11 → SCR-05 loan flow)\n- Quota expiry scheduler (60s cron)'),
        ('Sprint 4',
         'SCR-11 Portal, Trade-in, CRM, Refund',
         '- SCR-11 Cổng đơn hàng đầy đủ\n- Trade-in Platform (submit → inspect → ACCEPTED/FAILED)\n- Auto-update final_price khi trade-in ACCEPTED\n- SCR-06 CRM Leads Kanban + Quick-action bar\n- SCR-07 Discount approval (3-level: Sale → Manager → Admin)\n- Refund flow 3 cấp + Manager Override\n- Refund SLA cron (daily escalation)\n- Notification Matrix implementation (all 23 events)\n- SCR-05 Dashboard KPI + reports\n- SCR-09-CONF Cấu hình hệ thống (system_configs)'),
        ('Sprint 5',
         'Security Audit, Performance, E2E Testing',
         '- Mock KMS Decrypt Audit Logging (SCR-13)\n- Email verification flow complete\n- CORS policy config + security headers\n- audit_logs partitioning (monthly, archive >12M)\n- k6 Load Test: 500 VU × 5 phút, P99 < 200ms\n- Playwright E2E: 12 critical test scenarios\n- Newman CLI: 250+ API test cases\n- Lighthouse CI integration\n- Credit expiry cron (12 tháng)\n- Dark mode (enhancement backlog)\n- DB-level COMMENT ON TABLE/COLUMN'),
    ]
)

# Final signature
doc.add_paragraph()
final_para = doc.add_paragraph()
final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = final_para.add_run(
    '═══════════════════════════════════════════════════════════════════\n'
    'AutoDealership Enterprise Platform — Specification Document v24.0\n'
    'Ngày cập nhật: 2026-08-23  |  Áp dụng 152 findings từ Review Report\n'
    '✅ CRITICAL: 18/18  ✅ HIGH: 42/42  ✅ MEDIUM: 56/56  ✅ LOW: 36/36\n'
    'Style: Times New Roman — Unified Design System\n'
    '═══════════════════════════════════════════════════════════════════'
)
fr.font.name = FONT_CODE
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

# Save
OUT_PATH = 'Tai_Lieu_Dac_Ta_He_Thong_AutoDealership_v24_Final.docx'
doc.save(OUT_PATH)
print(f'✅ Done! Saved: {OUT_PATH}')
import os
size = os.path.getsize(OUT_PATH)
print(f'   File size: {size:,} bytes ({size//1024} KB)')
